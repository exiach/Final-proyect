"""Refina las figuras académicas y su ubicación en el documento final.

La edición conserva el archivo de entrada y produce una copia nueva. Las
figuras se construyen únicamente con datos, rutas y reglas presentes en el
repositorio; no se generan resultados experimentales nuevos.
"""

from __future__ import annotations

import os
import subprocess
from pathlib import Path

from PIL import Image, ImageFilter
from reportlab.lib.colors import HexColor, white
from reportlab.pdfgen import canvas
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "documentacion/Proyecto_Revisado_Academicamente_Ampliado.docx"
OUTPUT = ROOT / "documentacion/Proyecto_Revisado_Academicamente_Visuales.docx"
FIG_DIR = ROOT / "documentacion/figuras"

BLUE = "#1F4E79"
BLUE_2 = "#2F75B5"
BLUE_PALE = "#EAF2F8"
INK = "#1F2937"
MUTED = "#52606D"
LINE = "#7B8794"
AMBER = "#B7791F"
AMBER_PALE = "#FFF4D6"
GREEN = "#2E7D32"
GREEN_PALE = "#E8F5E9"
RED = "#B42318"
RED_PALE = "#FDECEC"
GRAY_PALE = "#F3F4F6"


def _pdf_to_png(pdf_path: Path, output: Path) -> Path:
    prefix = output.with_suffix("")
    subprocess.run(
        ["pdftoppm", "-png", "-singlefile", "-r", "300", str(pdf_path), str(prefix)],
        check=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
    )
    pdf_path.unlink(missing_ok=True)
    return output


def _box(c, x, y, w, h, face, edge, radius=8, width=1.4):
    c.setFillColor(HexColor(face))
    c.setStrokeColor(HexColor(edge))
    c.setLineWidth(width)
    c.roundRect(x, y, w, h, radius, stroke=1, fill=1)


def _text(c, x, y, value, size=10, color=INK, bold=False, align="center", leading=None):
    font = "Helvetica-Bold" if bold else "Helvetica"
    c.setFont(font, size)
    c.setFillColor(HexColor(color))
    lines = value.split("\n")
    leading = leading or size * 1.25
    for index, line in enumerate(lines):
        yy = y - index * leading
        if align == "left":
            c.drawString(x, yy, line)
        elif align == "right":
            c.drawRightString(x, yy, line)
        else:
            c.drawCentredString(x, yy, line)


def _arrow(c, x1, y1, x2, y2, color=LINE, width=1.5, dashed=False):
    c.setStrokeColor(HexColor(color))
    c.setFillColor(HexColor(color))
    c.setLineWidth(width)
    c.setDash(5, 4) if dashed else c.setDash()
    c.line(x1, y1, x2, y2)
    import math
    angle = math.atan2(y2 - y1, x2 - x1)
    length = 8
    spread = 0.48
    points = [
        (x2, y2),
        (x2 - length * math.cos(angle - spread), y2 - length * math.sin(angle - spread)),
        (x2 - length * math.cos(angle + spread), y2 - length * math.sin(angle + spread)),
    ]
    path = c.beginPath()
    path.moveTo(*points[0])
    path.lineTo(*points[1])
    path.lineTo(*points[2])
    path.close()
    c.drawPath(path, stroke=0, fill=1)
    c.setDash()


def _diamond(c, cx, cy, w, h, text):
    c.setFillColor(HexColor(AMBER_PALE))
    c.setStrokeColor(HexColor(AMBER))
    c.setLineWidth(1.4)
    path = c.beginPath()
    path.moveTo(cx, cy + h / 2)
    path.lineTo(cx + w / 2, cy)
    path.lineTo(cx, cy - h / 2)
    path.lineTo(cx - w / 2, cy)
    path.close()
    c.drawPath(path, stroke=1, fill=1)
    lines = text.split("\n")
    start = cy + (len(lines) - 1) * 6
    _text(c, cx, start, text, size=9.3, bold=True, leading=12)


def _phase(c, x, y, number, title, action, evidence, oe):
    _box(c, x, y, 292, 102, "#FFFFFF", BLUE_2, radius=8, width=1.5)
    _box(c, x + 12, y + 68, 34, 24, BLUE, BLUE, radius=10, width=1)
    _text(c, x + 29, y + 76, str(number), size=10, color="#FFFFFF", bold=True)
    _text(c, x + 55, y + 77, title, size=11.2, color=BLUE, bold=True, align="left")
    _text(c, x + 14, y + 47, action, size=9.0, align="left")
    _text(c, x + 14, y + 24, evidence, size=8.2, color=MUTED, align="left")
    _box(c, x + 240, y + 10, 39, 20, BLUE_PALE, BLUE_2, radius=8, width=0.8)
    _text(c, x + 259.5, y + 16, oe, size=7.4, color=BLUE, bold=True)


def generate_crisp_dm() -> Path:
    out = FIG_DIR / "fig_3_2_crisp_dm_academico.png"
    pdf = out.with_suffix(".pdf")
    c = canvas.Canvas(str(pdf), pagesize=(720, 680))
    c.setFillColor(white)
    c.rect(0, 0, 720, 680, fill=1, stroke=0)
    phases = [
        (38, 520, 1, "Comprensión del problema", "Definir rezago, alcance y uso docente", "Salida: problema y objetivos verificables", "OE1–5"),
        (390, 520, 2, "Comprensión de los datos", "Revisar 36 boletines de 2022–2024", "Salida: 1.118 observaciones estudiante-año", "OE1"),
        (390, 325, 3, "Preparación de los datos", "Validar, normalizar y construir T → T+1", "Salida: muestra longitudinal y control de faltantes", "OE1–3"),
        (38, 325, 4, "Modelado", "Entrenar Árbol, Random Forest y MLP", "Salida: modelos, escalador e hiperparámetros", "OE3"),
        (38, 130, 5, "Evaluación", "Aplicar prueba temporal 2023 → 2024", "Salida: matrices y métricas comparables", "OE4"),
        (390, 130, 6, "Despliegue", "Integrar Streamlit y resguardo pedagógico", "Salida: prototipo de apoyo a la decisión", "OE5"),
    ]
    for item in phases:
        _phase(c, *item)
    _arrow(c, 334, 571, 384, 571)
    _arrow(c, 536, 512, 536, 433)
    _arrow(c, 384, 376, 340, 376)
    _arrow(c, 184, 317, 184, 238)
    _arrow(c, 334, 181, 384, 181)
    _arrow(c, 536, 122, 536, 92, color=BLUE_2)
    c.setStrokeColor(HexColor(BLUE_2)); c.setLineWidth(1.2); c.setDash(5, 4)
    c.line(24, 170, 24, 376); c.line(24, 376, 34, 376)
    c.line(696, 170, 696, 571); c.line(696, 571, 686, 571)
    c.setDash()
    _text(c, 31, 274, "ajuste", size=7.8, color=BLUE, bold=True, align="left")
    _text(c, 688, 370, "seguimiento", size=7.8, color=BLUE, bold=True, align="right")
    _box(c, 90, 40, 540, 34, GRAY_PALE, "#D1D5DB", radius=7, width=0.7)
    _text(c, 360, 52, "Proceso iterativo: evaluación y uso retroalimentan la preparación y el modelado.", size=8.4, color=MUTED)
    c.showPage(); c.save()
    return _pdf_to_png(pdf, out)


def generate_architecture() -> Path:
    out = FIG_DIR / "fig_3_7_arquitectura_modular_academica.png"
    pdf = out.with_suffix(".pdf")
    c = canvas.Canvas(str(pdf), pagesize=(720, 300))
    c.setFillColor(white); c.rect(0, 0, 720, 300, fill=1, stroke=0)
    cards = [
        (18, "1", "Datos", "primaria_dataset.csv\n1.118 registros", "#F3E8FF", "#7E57A1"),
        (194, "2", "Modelos", "Random Forest · MLP\nStandardScaler", BLUE_PALE, BLUE_2),
        (370, "3", "Motor híbrido", "src/predictor.py\ninferencia + resguardo", AMBER_PALE, AMBER),
        (546, "4", "Interfaz", "Streamlit: monitoreo\nficha y simulador", GREEN_PALE, GREEN),
    ]
    for x, number, title, detail, face, edge in cards:
        _box(c, x, 62, 156, 194, face, edge, radius=8, width=1.4)
        _box(c, x + 10, 222, 26, 22, edge, edge, radius=8, width=0.8)
        _text(c, x + 23, 229, number, size=8.8, color="#FFFFFF", bold=True)
        _text(c, x + 44, 229, title, size=9.5, color=edge, bold=True, align="left")
        _text(c, x + 78, 155, detail, size=8.6, leading=14)
    for x in (176, 352, 528):
        _arrow(c, x, 158, x + 16, 158, width=1.7)
    _text(c, 360, 28, "Datos validados → inferencia → regla operativa → visualización para revisión docente", size=8.4, color=MUTED)
    c.showPage(); c.save()
    return _pdf_to_png(pdf, out)


def generate_safeguard() -> Path:
    out = FIG_DIR / "fig_3_8_resguardo_pedagogico_academico.png"
    pdf = out.with_suffix(".pdf")
    c = canvas.Canvas(str(pdf), pagesize=(720, 790))
    c.setFillColor(white); c.rect(0, 0, 720, 790, fill=1, stroke=0)
    _box(c, 190, 716, 340, 42, BLUE_PALE, BLUE_2, radius=12)
    _text(c, 360, 732, "Entrada: promedio general y materias reprobadas", size=10, color=BLUE, bold=True)
    _diamond(c, 360, 662, 260, 72, "¿Las nueve calificaciones\nestán completas?")
    _arrow(c, 360, 714, 360, 699)
    _box(c, 25, 600, 150, 55, GRAY_PALE, LINE, radius=8)
    _text(c, 100, 632, "SIN DATOS", size=10, color=MUTED, bold=True)
    _text(c, 100, 614, "No se calcula alerta", size=8.2, color=MUTED)
    _arrow(c, 230, 662, 176, 628); _text(c, 204, 650, "No", size=8, color=MUTED, bold=True)
    _box(c, 205, 568, 310, 58, BLUE_PALE, BLUE_2, radius=8)
    _text(c, 360, 605, "Inferencia del modelo", size=10, color=BLUE, bold=True)
    _text(c, 360, 584, "p_modelo se conserva como salida estadística", size=8.2, color=MUTED)
    _arrow(c, 360, 626, 360, 628); _text(c, 376, 640, "Sí", size=8, color=BLUE, bold=True)
    _diamond(c, 360, 500, 300, 80, "¿Promedio < 51\no reprobadas ≥ 2?")
    _arrow(c, 360, 566, 360, 542)
    _box(c, 22, 452, 205, 58, RED_PALE, RED, radius=8)
    _text(c, 124, 490, "Regla de resguardo", size=9.2, color=RED, bold=True)
    _text(c, 124, 469, "p_operativa = max(p_modelo, 0,85)", size=8.0)
    _arrow(c, 210, 500, 228, 485, color=RED); _text(c, 205, 510, "Sí", size=8, color=RED, bold=True)
    _diamond(c, 360, 392, 320, 82, "¿Una reprobada o\n51 ≤ promedio < 60?")
    _arrow(c, 360, 458, 360, 435); _text(c, 376, 448, "No", size=8, color=MUTED, bold=True)
    _box(c, 22, 315, 205, 60, AMBER_PALE, AMBER, radius=8)
    _text(c, 124, 353, "Regla preventiva", size=9.2, color=AMBER, bold=True)
    _text(c, 124, 332, "p_operativa = max(p_modelo, 0,50)", size=8.0)
    _arrow(c, 200, 392, 228, 350, color=AMBER); _text(c, 205, 395, "Sí", size=8, color=AMBER, bold=True)
    _box(c, 493, 315, 205, 60, GREEN_PALE, GREEN, radius=8)
    _text(c, 595, 353, "Sin ajuste pedagógico", size=9.0, color=GREEN, bold=True)
    _text(c, 595, 332, "p_operativa = p_modelo", size=8.2)
    _arrow(c, 520, 392, 492, 350, color=GREEN); _text(c, 505, 395, "No", size=8, color=GREEN, bold=True)
    _box(c, 120, 218, 480, 56, "#FFFFFF", BLUE, radius=8, width=1.4)
    _text(c, 360, 253, "Categorización común por umbrales operativos", size=9.4, color=BLUE, bold=True)
    _text(c, 360, 232, "Alto: p ≥ 0,70 · Medio: 0,40 ≤ p < 0,70 · Bajo: p < 0,40", size=8.2)
    _arrow(c, 124, 314, 260, 275, width=1.1)
    _arrow(c, 595, 314, 460, 275, width=1.1)
    _arrow(c, 360, 350, 360, 275, width=1.1)
    results = [
        (35, RED_PALE, RED, "ALTO", "Revisión prioritaria"),
        (260, AMBER_PALE, AMBER, "MEDIO", "Seguimiento preventivo"),
        (485, GREEN_PALE, GREEN, "BAJO", "Acompañamiento estándar"),
    ]
    for x, face, edge, title, subtitle in results:
        _box(c, x, 122, 200, 56, face, edge, radius=8)
        _text(c, x + 100, 156, title, size=9.5, color=edge, bold=True)
        _text(c, x + 100, 136, subtitle, size=7.9)
    _arrow(c, 360, 216, 135, 180, color=RED, width=1.0)
    _arrow(c, 360, 216, 360, 180, color=AMBER, width=1.0)
    _arrow(c, 360, 216, 585, 180, color=GREEN, width=1.0)
    _box(c, 65, 46, 590, 42, GRAY_PALE, "#D1D5DB", radius=7, width=0.7)
    _text(c, 360, 62, "La regla modifica el puntaje operativo; no valida el modelo ni elimina falsos negativos futuros.", size=8.1, color=MUTED)
    c.showPage(); c.save()
    return _pdf_to_png(pdf, out)


def crop_application_screenshot(source: Path, output: Path) -> Path:
    """Recorta la interfaz principal, sin inventar ni alterar su contenido."""
    image = Image.open(source).convert("RGB")
    cropped = image.crop((300, 45, 1438, 900))
    cropped = cropped.filter(ImageFilter.UnsharpMask(radius=1.1, percent=115, threshold=3))
    cropped.save(output, quality=96)
    return output


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def set_keep_with_next(paragraph, value=True) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    keep = p_pr.find(qn("w:keepNext"))
    if keep is None:
        keep = OxmlElement("w:keepNext")
        p_pr.append(keep)
    keep.set(qn("w:val"), "1" if value else "0")


def style_figure_paragraph(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(3)
    set_keep_with_next(paragraph)


def style_caption(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(1)
    set_keep_with_next(paragraph)
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.size = Pt(10)
        run.italic = True


def style_source(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(7)
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.size = Pt(9)


def style_body(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.15
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.size = Pt(11)


def make_picture_paragraph(doc, path: Path, width=Inches(6.25)):
    paragraph = doc.add_paragraph()
    paragraph.add_run().add_picture(str(path), width=width)
    style_figure_paragraph(paragraph)
    return paragraph


def make_text_paragraph(doc, text: str, kind="body"):
    paragraph = doc.add_paragraph()
    paragraph.add_run(text)
    if kind == "caption":
        style_caption(paragraph)
    elif kind == "source":
        style_source(paragraph)
    else:
        style_body(paragraph)
    return paragraph


def insert_after(anchor, paragraphs) -> None:
    current = anchor._p
    for paragraph in paragraphs:
        current.addnext(paragraph._p)
        current = paragraph._p


def find_exact(doc, text: str):
    return next(p for p in doc.paragraphs if p.text.strip() == text)


def find_starts(doc, prefix: str):
    return next(p for p in doc.paragraphs if p.text.strip().startswith(prefix))


def paragraph_index(paragraphs, target) -> int:
    return next(i for i, paragraph in enumerate(paragraphs) if paragraph._p is target._p)


def find_body_caption(doc, prefix: str):
    paragraphs = list(doc.paragraphs)
    for candidate in paragraphs:
        if not candidate.text.strip().startswith(prefix):
            continue
        idx = paragraph_index(paragraphs, candidate)
        if any(
            p._p.xpath(".//a:blip") or p._p.xpath(".//*[local-name()='imagedata']")
            for p in paragraphs[max(0, idx - 4):idx]
        ):
            return candidate
    raise ValueError(f"No se encontró el pie corporal: {prefix}")


def replace_inline_picture(paragraph, image_path: Path, width=Inches(6.15)) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)
    paragraph.add_run().add_picture(str(image_path), width=width)
    style_figure_paragraph(paragraph)


def remove_figure_by_caption(doc, prefix: str) -> None:
    paragraphs = list(doc.paragraphs)
    caption = None
    image_p = None
    idx = -1
    for candidate in paragraphs:
        if not candidate.text.strip().startswith(prefix):
            continue
        candidate_idx = paragraph_index(paragraphs, candidate)
        nearby = next(
            (
                p for p in reversed(paragraphs[max(0, candidate_idx - 4):candidate_idx])
                if p._p.xpath(".//a:blip") or p._p.xpath(".//*[local-name()='imagedata']")
            ),
            None,
        )
        if nearby is not None:
            caption, image_p, idx = candidate, nearby, candidate_idx
            break
    if caption is None or image_p is None:
        raise ValueError(f"No se encontró la figura corporal para: {prefix}")
    following = paragraphs[idx + 1] if idx + 1 < len(paragraphs) else None
    remove_paragraph(image_p)
    remove_paragraph(caption)
    if following is not None and following.text.strip().startswith("Fuente:"):
        remove_paragraph(following)


def remove_list_entries(doc, prefixes: tuple[str, ...]) -> None:
    paragraphs = list(doc.paragraphs)
    start = next(i for i, p in enumerate(paragraphs) if p.text.strip() == "Lista de figuras")
    end = next(i for i, p in enumerate(paragraphs) if p.text.strip() == "Lista de tablas")
    for paragraph in list(paragraphs[start + 1:end]):
        if paragraph.text.strip().startswith(prefixes):
            remove_paragraph(paragraph)


def remove_next_image_after(doc, paragraph_prefix: str, window: int = 5) -> None:
    paragraphs = list(doc.paragraphs)
    anchor = next(p for p in paragraphs if p.text.strip().startswith(paragraph_prefix))
    idx = paragraph_index(paragraphs, anchor)
    image_p = next(
        (
            p for p in paragraphs[idx + 1:idx + 1 + window]
            if p._p.xpath(".//a:blip") or p._p.xpath(".//*[local-name()='imagedata']")
        ),
        None,
    )
    if image_p is None:
        raise ValueError(f"No se encontró una captura después de: {paragraph_prefix}")
    image_idx = paragraph_index(paragraphs, image_p)
    following = paragraphs[image_idx + 1] if image_idx + 1 < len(paragraphs) else None
    remove_paragraph(image_p)
    if (
        following is not None
        and not following.text.strip()
        and following._p.find("./w:pPr/w:sectPr", following._p.nsmap) is None
    ):
        remove_paragraph(following)


def insert_figure_list_entries(doc, entries: list[str]) -> None:
    anchor = find_exact(doc, "Lista de tablas")
    for text in entries:
        paragraph = doc.add_paragraph()
        paragraph.add_run(text)
        paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(6.5))
        for run in paragraph.runs:
            run.font.name = "Arial"
            run.font.size = Pt(11)
        anchor._p.addprevious(paragraph._p)


def main() -> None:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    crisp = FIG_DIR / "fig_3_2_crisp_dm_academico.png"
    architecture = FIG_DIR / "fig_3_7_arquitectura_modular_academica.png"
    safeguard = FIG_DIR / "fig_3_8_resguardo_pedagogico_academico.png"
    monitoring = FIG_DIR / "fig_4_8_prototipo_monitoreo_academico.png"
    simulator = FIG_DIR / "fig_4_9_prototipo_simulador_academico.png"
    if not crisp.exists():
        generate_crisp_dm()
    if not architecture.exists():
        generate_architecture()
    if not safeguard.exists():
        generate_safeguard()
    if not monitoring.exists():
        crop_application_screenshot(FIG_DIR / "fig_4_8_prototipo_monitoreo_curso.png", monitoring)
    if not simulator.exists():
        crop_application_screenshot(FIG_DIR / "fig_4_10_prototipo_simulador_libre.png", simulator)

    doc = Document(SOURCE)

    # Figura 3.2: imagen, referencia cruzada y fuente.
    crisp_caption = find_body_caption(doc, "Figura 3.2. Flujo metodológico CRISP-DM adaptado")
    crisp_paragraphs = list(doc.paragraphs)
    crisp_image = crisp_paragraphs[paragraph_index(crisp_paragraphs, crisp_caption) - 1]
    replace_inline_picture(crisp_image, crisp)
    methodology_lead = find_starts(doc, "En la Figura 3-2")
    methodology_lead.text = (
        "La Figura 3.2 resume la adaptación de CRISP-DM al proyecto y vincula cada fase con sus productos verificables: "
        "consolidación de datos, construcción longitudinal, entrenamiento, evaluación temporal y despliegue del prototipo."
    )
    style_body(methodology_lead)
    source_crisp = find_exact(doc, "Fuente: elaboración propia (2026).")
    source_crisp.text = "Fuente: elaboración propia con base en Wirth y Hipp (2000) y los artefactos del proyecto (2026)."
    style_source(source_crisp)

    # Retira las composiciones de código. Las tablas y explicaciones metodológicas permanecen.
    for caption_prefix in (
        "Figura 3.7. Validación, construcción de variables",
        "Figura 3.8. Código reproducible",
        "Figura 3.9. Entrenamiento reproducible",
        "Figura 3.10. Trazabilidad de la segregación",
    ):
        remove_figure_by_caption(doc, caption_prefix)

    trace_intro = find_starts(doc, "Para fortalecer la reproducibilidad del capítulo")
    trace_intro.text = (
        "Para fortalecer la reproducibilidad del capítulo, esta sección relaciona cada objetivo específico con su entrada, "
        "procedimiento, control y evidencia verificable. Se priorizan tablas metodológicas, diagramas y salidas agregadas; "
        "las capturas de código se omiten porque el repositorio y los notebooks constituyen la fuente técnica consultable."
    )
    style_body(trace_intro)

    # Retira las cuatro capturas oscuras de código que permanecían en 3.4–3.5.
    # La explicación metodológica se conserva y la implementación queda verificable en los notebooks.
    for prefix in (
        "Para automatizar la consolidación de los boletines heterogéneos",
        "Tras procesar 36 boletines de 2022 a 2024",
        "Se evaluó la tasa de reprobación porcentual",
        "La comparación descriptiva mostró promedios menores",
    ):
        remove_next_image_after(doc, prefix)

    # Arquitectura en 3.9.1, donde corresponde metodológicamente.
    architecture_anchor = find_starts(doc, "src/ui/: Componentes modulares")
    insert_after(architecture_anchor, [
        make_picture_paragraph(doc, architecture),
        make_text_paragraph(doc, "Figura 3.7. Arquitectura modular y flujo de ejecución del prototipo", "caption"),
        make_text_paragraph(doc, "Fuente: elaboración propia con base en app.py, config.py y los módulos src/ (2026).", "source"),
    ])

    # Sustituye la figura bajo 3.9.2 y añade pie y fuente académicos.
    safeguard_heading = find_exact(doc, "3.9.2 Capa de Resguardo Pedagógico (Sistema Híbrido)")
    safeguard_paragraphs = list(doc.paragraphs)
    safeguard_text = next(
        p for p in safeguard_paragraphs[paragraph_index(safeguard_paragraphs, safeguard_heading) + 1:]
        if p.text.strip()
    )
    paragraphs = list(doc.paragraphs)
    idx = paragraph_index(paragraphs, safeguard_text)
    old_guard = next(
        (p for p in paragraphs[idx + 1:idx + 5] if p._p.xpath(".//a:blip") or p._p.xpath(".//*[local-name()='imagedata']")),
        None,
    )
    if old_guard is not None:
        remove_paragraph(old_guard)
    insert_after(safeguard_text, [
        make_picture_paragraph(doc, safeguard, width=Inches(6.05)),
        make_text_paragraph(doc, "Figura 3.8. Lógica de inferencia y resguardo pedagógico del sistema híbrido", "caption"),
        make_text_paragraph(doc, "Fuente: elaboración propia con base en src/predictor.py y config.py (2026).", "source"),
    ])

    # Renumera las capturas de interfaz que sí aportan evidencia en el capítulo 3.
    panel_caption = find_body_caption(doc, "Figura 3.11. Panel principal")
    panel_caption.text = (
        "Figura 3.9. Panel principal del prototipo con filtros e indicadores agregados"
    )
    style_caption(panel_caption)
    simulator_caption = find_body_caption(doc, "Figura 3.12. Simulador libre")
    simulator_caption.text = (
        "Figura 3.10. Simulador libre y separación entre probabilidad del modelo y puntaje operativo"
    )
    style_caption(simulator_caption)

    # 4.5: elimina el diagrama de arquitectura mal ubicado e incorpora el prototipo real.
    prototype_heading = find_exact(doc, "4.5 Resultados del Prototipo Docente de Alertas Tempranas (OE5)")
    prototype_paragraphs = list(doc.paragraphs)
    prototype_lead = next(
        p for p in prototype_paragraphs[paragraph_index(prototype_paragraphs, prototype_heading) + 1:]
        if p.text.strip()
    )
    paragraphs = list(doc.paragraphs)
    lead_idx = paragraph_index(paragraphs, prototype_lead)
    wrong_image = next(
        (p for p in paragraphs[lead_idx + 1:lead_idx + 4] if p._p.xpath(".//a:blip") or p._p.xpath(".//*[local-name()='imagedata']")),
        None,
    )
    if wrong_image is not None:
        remove_paragraph(wrong_image)
    insert_after(prototype_lead, [
        make_picture_paragraph(doc, monitoring),
        make_text_paragraph(doc, "Figura 4.8. Vista de monitoreo del curso y panel de alertas tempranas", "caption"),
        make_text_paragraph(doc, "Fuente: captura de la aplicación Streamlit desarrollada y ejecutada localmente (2026).", "source"),
        make_text_paragraph(
            doc,
            "La vista de monitoreo reúne los filtros de cohorte, el modelo seleccionado y los indicadores agregados. "
            "Su función es priorizar registros para revisión docente, no emitir decisiones automáticas.",
        ),
        make_picture_paragraph(doc, simulator),
        make_text_paragraph(doc, "Figura 4.9. Simulador libre para comprobar escenarios y reglas operativas", "caption"),
        make_text_paragraph(doc, "Fuente: captura de la aplicación Streamlit desarrollada y ejecutada localmente (2026).", "source"),
        make_text_paragraph(
            doc,
            "El simulador permite variar el promedio general y el número de materias reprobadas sin modificar el dataset. "
            "La interfaz muestra por separado la probabilidad del modelo y el puntaje operativo, lo que hace visible el efecto de la regla pedagógica.",
        ),
    ])

    # Sincroniza la lista de figuras con las eliminaciones y adiciones.
    remove_list_entries(doc, ("Figura 3.7.", "Figura 3.8.", "Figura 3.9.", "Figura 3.10.", "Figura 3.11.", "Figura 3.12.", "Figura 4.8.", "Figura 4.9."))
    insert_figure_list_entries(doc, [
        "Figura 3.7. Arquitectura modular y flujo de ejecución del prototipo\t—",
        "Figura 3.8. Lógica de inferencia y resguardo pedagógico del sistema híbrido\t—",
        "Figura 3.9. Panel principal del prototipo con filtros e indicadores agregados\t—",
        "Figura 3.10. Simulador libre y separación entre probabilidad y puntaje operativo\t—",
        "Figura 4.8. Vista de monitoreo del curso y panel de alertas tempranas\t—",
        "Figura 4.9. Simulador libre para comprobar escenarios y reglas operativas\t—",
    ])

    # Solicita actualizar campos al abrir el archivo en Word/Google Docs.
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
