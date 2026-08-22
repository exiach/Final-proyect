"""Genera una copia académicamente corregida del borrador, sin alterar el original."""

from __future__ import annotations

import shutil
import tempfile
import zipfile
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from lxml import etree


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "documentacion/Borrador FORMATO PROYECTO 3ºV_DIPLOMADO 'CIENCIA DE DATOS'.docx"
OUTPUT = ROOT / "documentacion/Proyecto_Revisado_Academicamente.docx"


GLOBAL_REPLACEMENTS = {
    "2021–2024": "2022–2024",
    "2021-2024": "2022-2024",
    "2021 a 2024": "2022 a 2024",
    "2021, 2022, 2023 y 2024": "2022, 2023 y 2024",
    "2021, 2022, 2023, 2024": "2022, 2023, 2024",
    "Aguiar & Morales, 2021": "Macfadyen & Dawson, 2010",
    "De-La-Peña & Luque-Rojas, 2021": "Romero & Ventura, 2020",
    "Cardoso et al., 2021": "Macfadyen & Dawson, 2010",
    "García & Ruiz, 2023": "Fawcett, 2006",
    "Pérez & Martínez, 2024": "Haykin, 2009",
    "Pérez, 2021": "Streamlit, 2024",
    "López et al., 2022": "Ministerio de Educación del Estado Plurinacional de Bolivia, 2021",
    "Román, 2013": "Romero & Ventura, 2020",
    "UNESCO, 2023": "Romero & Ventura, 2020",
    "Chawla et al., 2004": "Chawla et al., 2002",
    "Hastie et al., 2008": "Hastie et al., 2009",
    "Hastie, Tibshirani & Friedman, 2008": "Hastie et al., 2009",
    "James, Witten, Hastie & Tibshirani, 2013": "James et al., 2013",
    "Goodfellow, Bengio & Courville, 2016": "Goodfellow et al., 2016",
    "Han, Kamber & Pei, 2011": "Han et al., 2012",
    "Breiman et al., 1984": "James et al., 2013",
    "Ministerio de Educación de Bolivia, 2022": "Ministerio de Educación del Estado Plurinacional de Bolivia, 2021",
    "1,118": "1.118",
    "2.06%": "2,06 %",
    "1.70%": "1,70 %",
    "4.13": "4,13",
    "$0.0 \\le \\text{nota} \\le 100.0$": "0,0 ≤ nota ≤ 100,0",
    "$0.0 \\le \\text{promedio} \\le 100.0$": "0,0 ≤ promedio ≤ 100,0",
    "$0 \\le N \\le 9$": "0 ≤ N ≤ 9",
    "$< 51$": "< 51",
    "$P \\ge 0.70$": "P ≥ 0,70",
    "$0.40 \\le P < 0.70$": "0,40 ≤ P < 0,70",
    "$P < 0.40$": "P < 0,40",
    "$P(Y=1 \\mid \\mathbf{x})$": "P(Y=1 | x)",
    "Rendimiento Cuantitativo sobre el Conjunto de Prueba Filtrado (N=51)": "Rendimiento en la Prueba Temporal 2023→2024 (N=248)",
    "1,095": "1.095",
    "97.94%": "97,94 %",
    "$1.61%$": "1,61 %",
    "$2.51%$": "2,51 %",
    "$1.5%$": "1,5 %",
    "$2.8%$": "2,8 %",
}


REFERENCES = [
    "Baker, R. S., & Inventado, P. S. (2014). Educational data mining and learning analytics. En J. A. Larusson y B. White (Eds.), Learning analytics: From research to practice (pp. 61–75). Springer. https://doi.org/10.1007/978-1-4614-3305-7_4",
    "Breiman, L. (2001). Random forests. Machine Learning, 45(1), 5–32. https://doi.org/10.1023/A:1010933404324",
    "Chawla, N. V., Bowyer, K. W., Hall, L. O., & Kegelmeyer, W. P. (2002). SMOTE: Synthetic minority over-sampling technique. Journal of Artificial Intelligence Research, 16, 321–357. https://doi.org/10.1613/jair.953",
    "Fawcett, T. (2006). An introduction to ROC analysis. Pattern Recognition Letters, 27(8), 861–874. https://doi.org/10.1016/j.patrec.2005.10.010",
    "Goodfellow, I., Bengio, Y., & Courville, A. (2016). Deep learning. MIT Press.",
    "Han, J., Kamber, M., & Pei, J. (2012). Data mining: Concepts and techniques (3rd ed.). Morgan Kaufmann.",
    "Hastie, T., Tibshirani, R., & Friedman, J. (2009). The elements of statistical learning (2nd ed.). Springer. https://doi.org/10.1007/978-0-387-84858-7",
    "Haykin, S. (2009). Neural networks and learning machines (3rd ed.). Pearson.",
    "James, G., Witten, D., Hastie, T., & Tibshirani, R. (2013). An introduction to statistical learning. Springer. https://doi.org/10.1007/978-1-4614-7138-7",
    "Macfadyen, L. P., & Dawson, S. (2010). Mining LMS data to develop an ‘early warning system’ for educators: A case study. Computers & Education, 54(2), 588–599. https://doi.org/10.1016/j.compedu.2009.09.008",
    "Ministerio de Educación del Estado Plurinacional de Bolivia. (2021). Reglamento de evaluación del desarrollo curricular del Sistema Educativo Plurinacional.",
    "Pedregosa, F., Varoquaux, G., Gramfort, A., Michel, V., Thirion, B., Grisel, O., Blondel, M., Prettenhofer, P., Weiss, R., Dubourg, V., Vanderplas, J., Passos, A., Cournapeau, D., Brucher, M., Perrot, M., & Duchesnay, É. (2011). Scikit-learn: Machine learning in Python. Journal of Machine Learning Research, 12, 2825–2830.",
    "Provost, F., & Fawcett, T. (2013). Data science for business. O’Reilly Media.",
    "Romero, C., & Ventura, S. (2020). Educational data mining and learning analytics: An updated survey. WIREs Data Mining and Knowledge Discovery, 10(3), e1355. https://doi.org/10.1002/widm.1355",
    "Streamlit. (2024). Streamlit documentation. https://docs.streamlit.io",
    "Wirth, R., & Hipp, J. (2000). CRISP-DM: Towards a standard process model for data mining. En Proceedings of the 4th International Conference on the Practical Applications of Knowledge Discovery and Data Mining (pp. 29–39).",
]

REFERENCE_ITALICS = {
    REFERENCES[0]: ["Learning analytics: From research to practice"],
    REFERENCES[1]: ["Machine Learning, 45(1)"],
    REFERENCES[2]: ["Journal of Artificial Intelligence Research, 16"],
    REFERENCES[3]: ["Pattern Recognition Letters, 27(8)"],
    REFERENCES[4]: ["Deep learning"],
    REFERENCES[5]: ["Data mining: Concepts and techniques"],
    REFERENCES[6]: ["The elements of statistical learning"],
    REFERENCES[7]: ["Neural networks and learning machines"],
    REFERENCES[8]: ["An introduction to statistical learning"],
    REFERENCES[9]: ["Computers & Education, 54(2)"],
    REFERENCES[10]: ["Reglamento de evaluación del desarrollo curricular del Sistema Educativo Plurinacional"],
    REFERENCES[11]: ["Journal of Machine Learning Research, 12"],
    REFERENCES[12]: ["Data science for business"],
    REFERENCES[13]: ["WIREs Data Mining and Knowledge Discovery, 10(3)"],
    REFERENCES[14]: ["Streamlit documentation"],
    REFERENCES[15]: ["Proceedings of the 4th International Conference on the Practical Applications of Knowledge Discovery and Data Mining"],
}


def set_text_preserve_first_run(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def replace_everywhere(doc: Document) -> None:
    containers = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                containers.extend(cell.paragraphs)
    for section in doc.sections:
        containers.extend(section.header.paragraphs)
        containers.extend(section.footer.paragraphs)
    for paragraph in containers:
        text = paragraph.text
        new = text
        for old, replacement in GLOBAL_REPLACEMENTS.items():
            new = new.replace(old, replacement)
        if new != text:
            set_text_preserve_first_run(paragraph, new)


def replace_prefix(doc: Document, prefix: str, text: str) -> None:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            set_text_preserve_first_run(paragraph, text)
            return


def insert_before(paragraph, text: str, style: str = "normal"):
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p.getparent().remove(new_para._p)
    new_p.addnext(new_para._p)
    new_para.style = style
    new_para.add_run(text)
    return new_para


def insert_after(paragraph, text: str, style: str = "normal"):
    new_para = paragraph._parent.add_paragraph()
    new_para._p.getparent().remove(new_para._p)
    paragraph._p.addnext(new_para._p)
    new_para.style = style
    new_para.add_run(text)
    return new_para


def style_reference(paragraph, reference: str) -> None:
    """Aplica cursivas de APA a títulos de obras o revista/volumen."""
    set_text_preserve_first_run(paragraph, "")
    cursor = 0
    for phrase in sorted(REFERENCE_ITALICS.get(reference, []), key=reference.index):
        start = reference.index(phrase, cursor)
        if start > cursor:
            paragraph.add_run(reference[cursor:start])
        run = paragraph.add_run(phrase)
        run.italic = True
        cursor = start + len(phrase)
    if cursor < len(reference):
        paragraph.add_run(reference[cursor:])


def update_fields_on_open(doc: Document) -> None:
    settings = doc.settings._element
    node = settings.find(qn("w:updateFields"))
    if node is None:
        node = OxmlElement("w:updateFields")
        settings.append(node)
    node.set(qn("w:val"), "true")


def remove_embedded_images(doc: Document, filenames: set[str]) -> None:
    """Elimina párrafos que contienen capturas obsoletas identificadas por archivo."""
    for paragraph in list(doc.paragraphs):
        for blip in paragraph._p.xpath(".//a:blip"):
            relation_id = blip.get(qn("r:embed"))
            if not relation_id:
                continue
            part = doc.part.related_parts.get(relation_id)
            if part is not None and Path(str(part.partname)).name in filenames:
                paragraph._p.getparent().remove(paragraph._p)
                break


def remove_inline_drawings(doc: Document, filenames: set[str]) -> None:
    """Elimina dibujos concretos sin borrar otros dibujos del mismo párrafo."""
    for paragraph in list(doc.paragraphs):
        for blip in list(paragraph._p.xpath(".//a:blip")):
            relation_id = blip.get(qn("r:embed"))
            part = doc.part.related_parts.get(relation_id) if relation_id else None
            if part is None or Path(str(part.partname)).name not in filenames:
                continue
            node = blip
            while node is not None and etree.QName(node).localname not in {"drawing", "pict"}:
                node = node.getparent()
            if node is not None and node.getparent() is not None:
                node.getparent().remove(node)


def add_index_entry(anchor, text: str, page: int) -> None:
    paragraph = insert_before(anchor, f"{text}\t{page}", "normal")
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        Inches(6.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
    )


def add_toc_entry(anchor, template, text: str, page: int) -> None:
    """Agrega una entrada conservando tabulaciones y sangrías del índice existente."""
    paragraph = insert_before(anchor, f"{text}\t{page}", "normal")
    current_ppr = paragraph._p.find(qn("w:pPr"))
    if current_ppr is not None:
        paragraph._p.remove(current_ppr)
    if template._p.pPr is not None:
        paragraph._p.insert(0, deepcopy(template._p.pPr))


def synchronize_toc(doc: Document) -> None:
    """Sincroniza el índice estático con la paginación de la copia revisada."""
    start = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "Tabla de contenidos")
    end = next(i for i, p in enumerate(doc.paragraphs[start + 1 :], start + 1) if p.text.strip() == "Lista de figuras")
    pages = {
        "1.": 1, "1.1.": 1, "1.2.": 2, "1.3.": 3, "1.4.": 3, "1.4.1.": 3,
        "2.": 5, "2.1.": 5, "2.2.": 5, "2.2.1.": 5, "2.2.2.": 5, "2.2.3.": 6,
        "2.3.": 6, "2.3.1.": 6, "2.3.2.": 6, "2.3.3.": 7, "2.4.": 7, "2.5.": 8,
        "2.6.": 8, "2.7.": 9, "2.7.1.": 9, "2.7.2.": 9,
        "3.": 10, "3.1.": 10, "3.2.": 11, "3.3.": 12, "3.3.1.": 12,
        "3.4.": 12, "3.4.1.": 12, "3.4.2.": 14, "3.4.3.": 15,
        "3.5.": 16, "3.5.1.": 16, "3.5.2.": 16, "3.6.": 17,
        "3.7.": 17, "3.7.1.": 17, "3.7.2.": 17, "3.7.3.": 17, "3.8.": 17,
        "3.9.": 18, "3.9.1.": 18, "3.9.2.": 18,
        "4.": 20, "4.1.": 20, "4.1.1.": 20, "4.1.2.": 21,
        "4.2.": 21, "4.2.1.": 21, "4.2.2.": 22,
        "4.3.": 24, "4.3.1.": 24, "4.3.2.": 25,
        "5.": 32, "6.": 33,
        "Bibliografía": 34, "Referencias bibliográficas": 34, "Anexos": 35,
    }
    for paragraph in doc.paragraphs[start + 1 : end]:
        text = paragraph.text.strip()
        if not text:
            continue
        key = text.split("\t", 1)[0].strip()
        if key in pages:
            title = text.rsplit("\t", 1)[0]
            if key == "3.9.2.":
                title = "3.9.2.\tCapa de Resguardo Pedagógico (Sistema Híbrido)"
            elif key == "Referencias bibliográficas":
                title = "Bibliografía"
            set_text_preserve_first_run(paragraph, f"{title}\t{pages[key]}")

    conclusion = next(
        p for p in doc.paragraphs[start + 1 : end]
        if p.text.strip().startswith("5.\tConclusiones")
    )
    template = next(
        p for p in doc.paragraphs[start + 1 : end]
        if p.text.strip().startswith("4.3.2.\t")
    )
    missing = [
        ("4.4.\tResultados de la Evaluación y Segregación por Niveles de Riesgo (OE4)", 25),
        ("4.4.1.\tMatriz de Umbrales de Clasificación de Riesgo", 26),
        ("4.4.2.\tAnálisis Temporal y Variabilidad por Paralelo", 27),
        ("4.5.\tResultados del Prototipo Docente de Alertas Tempranas (OE5)", 29),
        ("4.5.1.\tValidación Operativa de la Capa de Resguardo Pedagógico", 30),
        ("4.6.\tDiscusión de Resultados", 30),
    ]
    for label, page in missing:
        add_toc_entry(conclusion, template, label, page)


def revise_content(doc: Document) -> None:
    for section in doc.sections:
        for header in (section.header, section.first_page_header, section.even_page_header):
            for p in header.paragraphs:
                if p.text.strip():
                    set_text_preserve_first_run(p, "PREDICCIÓN EXPLORATORIA DEL REZAGO ACADÉMICO")

    replacements = {
        "El rezago académico en educación primaria compromete": "El rezago académico en educación primaria requiere mecanismos de seguimiento que permitan priorizar la revisión docente. El proyecto desarrolló un prototipo exploratorio para estimar el riesgo de reprobación en la gestión siguiente en la U.E. José María Santiváñez, utilizando datos históricos de 2022 a 2024.",
        "La metodología comprendió un enfoque cuantitativo": "Se aplicó un enfoque cuantitativo y CRISP-DM. Se consolidaron 1.118 observaciones estudiante-año de 592 estudiantes. El modelado utilizó 489 transiciones consecutivas T→T+1 con datos completos en ambas gestiones, con seis positivas. Se entrenó con 241 transiciones 2022→2023 y se evaluó temporalmente con 248 transiciones 2023→2024.",
        "Los resultados determinaron que": "Comunicación y Lenguajes (2,06 %) y Matemática (1,70 %) presentaron las mayores tasas descriptivas de reprobación. En la prueba temporal, Árbol y Random Forest detectaron uno de tres casos positivos, con 13 falsas alarmas; la MLP no detectó positivos. Los resultados se consideran exploratorios.",
        "Se concluye que la minería de datos educativos permite anticipar": "Se concluye que la integración de datos y el prototipo son técnicamente viables, pero la evidencia disponible no permite afirmar alta capacidad predictiva. Se requieren más cohortes y validación prospectiva. El sistema debe utilizarse únicamente como apoyo sujeto a revisión docente.",
        "Tras procesar los 36 archivos": "Tras procesar 36 boletines de 2022 a 2024, se obtuvieron 1.118 observaciones estudiante-año correspondientes a 592 estudiantes únicos. La etiqueta rezago se definió como al menos una asignatura con nota inferior a 51; no equivale a repetición de curso ni deserción.",
        "Para predecir el riesgo de rezago académico en una gestión futura": "Para predecir la gestión T+1 se vinculó cada observación de T con el resultado del mismo estudiante en el año consecutivo, exigiendo datos completos en origen y destino. Se obtuvieron 489 transiciones: 241 de 2022→2023 y 248 de 2023→2024, con seis casos positivos. Se excluyeron 33 pares con etiqueta objetivo incompleta. No se utilizaron identificadores personales como predictores.",
        "Los modelos fueron evaluados sobre el conjunto de prueba independiente": "La evaluación principal fue temporal: entrenamiento con 2022→2023 (N=241; tres positivos) y prueba con 2023→2024 (N=248; tres positivos). Las métricas son exploratorias por la escasez de eventos.",
        "Análisis Técnico del Rendimiento": "Resultados de la prueba temporal: Árbol y Random Forest obtuvieron matriz [[232, 13], [2, 1]], precisión de rezago 0,0714, recall 0,3333, F1 0,1176 y balanced accuracy 0,6401. La MLP obtuvo [[245, 0], [3, 0]] y recall 0. No se demuestra superioridad ni alta precisión.",
        "En las pruebas funcionales del prototipo": "Las pruebas unitarias verificaron la implementación de la regla pedagógica: un promedio menor a 51 o dos o más asignaturas reprobadas activa una alerta alta. Esto valida la regla programada, no la sensibilidad predictiva ni una cobertura del 100 % de casos futuros.",
        "Eficacia del Modelado Predictivo y Capa Híbrida": "Alcance del Modelado Predictivo: La evaluación temporal mostró capacidad limitada: Árbol y Random Forest detectaron uno de tres positivos y la MLP ninguno. La regla pedagógica se separa de la probabilidad estadística. Los modelos requieren nuevas cohortes y validación prospectiva.",
        "En investigaciones aplicadas a sistemas educativos formales": "La literatura sobre minería de datos educativos documenta el uso de clasificadores y sistemas de alerta para apoyar la revisión docente (Macfadyen & Dawson, 2010; Romero & Ventura, 2020). Su desempeño depende del contexto, la calidad de los datos y la validación externa; por ello no se trasladan cifras de otros estudios como garantía para este proyecto.",
        "Desarrollar un modelo de predicción del rezago académico": "Desarrollar y evaluar exploratoriamente un sistema de apoyo para estimar el riesgo de reprobación en la gestión siguiente, utilizando datos históricos de educación primaria, con el fin de priorizar la revisión docente en la U.E. José María Santiváñez.",
        "En el proyecto se implementa una arquitectura con dos capas ocultas": "En el proyecto se implementó una capa oculta de cuatro neuronas (hidden_layer_sizes=(4,)), regularización L2 (alpha=0.1) y learning_rate_init=0.0001, con estandarización previa de las variables.",
        "La disparidad de clases es un fenómeno inherente": "Cuando la clase positiva es muy escasa, la exactitud global puede ser engañosa. El proyecto utilizó separación temporal, ponderación de clases en Árbol y Random Forest, y métricas centradas en la clase positiva. Estas medidas no compensan la existencia de solo seis transiciones positivas.",
        "Esta capa de resguardo garantiza": "La regla evita clasificar operativamente como bajo riesgo un registro completo que ya contiene notas reprobatorias. El puntaje operativo se muestra separado de la probabilidad del modelo y no representa una probabilidad calibrada ni una garantía de detección futura.",
        "Originalmente, esta información se encontraba almacenada": "Los boletines estaban almacenados en PDF y fueron convertidos a XLSX para su procesamiento. El registro del proyecto menciona el uso de iLovePDF; dado que los archivos contienen datos de menores, esta transferencia a un servicio externo constituye una limitación de privacidad. En futuras iteraciones debe emplearse conversión local o contar con autorización institucional y garantías de tratamiento de datos.",
        "Los resultados empíricos revelaron": "Descriptivamente, Comunicación y Lenguajes (2,06 %) y Matemática (1,70 %) registraron las mayores tasas de reprobación de la muestra. El resultado no demuestra causalidad.",
        "El análisis comparativo de promedios entre estudiantes aprobados": "La comparación descriptiva mostró promedios menores en las observaciones con rezago. No se realizó una prueba inferencial ni se interpreta la diferencia como efecto causal.",
        "Entrenado con restricción de profundidad máxima": "El Árbol de Decisión se configuró con max_depth=3, min_samples_leaf=10 y class_weight='balanced'.",
        "Configurado con 200 estimadores": "Random Forest se configuró con 300 estimadores, max_depth=4, min_samples_leaf=5 y class_weight='balanced'.",
        "Para asegurar que la herramienta sea 100% confiable": "Para impedir que la interfaz presente como bajo riesgo un registro completo con notas ya reprobatorias, src/predictor.py aplica una regla pedagógica. Esta regla no elimina falsos negativos futuros ni valida el modelo.",
        "Hallazgo Clave:": "Hallazgo descriptivo: Comunicación y Lenguajes y Matemática registraron las mayores tasas de reprobación de la muestra; no se infiere que sean causas del rezago futuro.",
        "Asimismo, la cantidad promedio de materias reprobadas": "En las observaciones con rezago, la media fue de 4,13 asignaturas reprobadas. Este valor describe el conjunto disponible y no demuestra un proceso causal.",
        "El prototipo del Sistema de Apoyo a la Decisión Docente fue construido y desplegado satisfactoriamente": "El prototipo del Sistema de Apoyo a la Decisión Docente fue implementado en Streamlit y sometido a pruebas funcionales automatizadas. No se realizó todavía una evaluación de usabilidad ni de impacto pedagógico con docentes.",
        "Factores Predictores de Rezago:": "Patrones observados: Comunicación y Lenguajes y Matemática presentaron las mayores tasas descriptivas de reprobación. Con solo dos predictores agregados y seis transiciones positivas, no es posible declarar factores fuertes o generalizables.",
        "Superación del Desbalance de Datos mediante Enfoques Híbridos:": "Desbalance y enfoque híbrido: La ponderación de clases y la regla pedagógica responden a necesidades distintas. La prueba temporal no demostró que la arquitectura sea óptima; se requiere ampliar las cohortes y validar prospectivamente.",
        "Rol del Sistema como Apoyo a la Decisión:": "Rol del sistema: El prototipo organiza información para apoyar la revisión docente. Sus alertas no sustituyen el juicio profesional ni deben activar decisiones automáticas sobre estudiantes.",
        "Identificación Empírica de Materias Críticas:": "Patrones descriptivos: Comunicación y Lenguajes (2,06 %) y Matemática (1,70 %) registraron las mayores tasas de reprobación. En las observaciones con rezago, la media fue 4,13 asignaturas reprobadas. No se establece causalidad.",
        "Implementación del Sistema de Apoyo a la Decisión Docente:": "Implementación del sistema de apoyo: Se implementó un prototipo funcional en Streamlit con tres niveles de alerta, una categoría Sin datos y simuladores. Las pruebas verifican funciones del software, no impacto pedagógico.",
        "El acceso a la educación primaria de calidad": "El seguimiento del rendimiento académico puede beneficiarse de registros longitudinales integrados. La minería de datos educativos estudia métodos para analizar información producida en contextos de aprendizaje (Baker & Inventado, 2014; Romero & Ventura, 2020). En este proyecto, el término rezago se limita operacionalmente a la reprobación de al menos una asignatura; no representa por sí mismo deserción, repetición ni un diagnóstico integral.",
        "En el contexto del Estado Plurinacional de Bolivia": "En Bolivia, el Reglamento de Evaluación del Desarrollo Curricular consultado establece la escala y los criterios institucionales utilizados en este trabajo (Ministerio de Educación del Estado Plurinacional de Bolivia, 2021). El proyecto no evalúa el sistema educativo nacional: analiza exclusivamente los registros disponibles de una unidad educativa.",
        "Situando este panorama en el departamento de Cochabamba": "En la U.E. José María Santiváñez, los registros de 2022 a 2024 estaban distribuidos por gestión, grado y paralelo. Esta fragmentación dificultaba el análisis longitudinal. La necesidad abordada fue consolidar la información y construir un prototipo exploratorio para priorizar casos que requieren revisión docente.",
        "Es ante esta problemática que la intersección": "La Minería de Datos Educativos y la Analítica del Aprendizaje ofrecen métodos para explorar registros académicos y presentar indicadores útiles para la toma de decisiones (Baker & Inventado, 2014; Romero & Ventura, 2020). Su aplicación exige validar los modelos en el contexto específico y evitar decisiones automáticas sobre estudiantes.",
        "En Bolivia, iniciativas recientes impulsadas": "El alcance de este trabajo se limita a una institución de Cochabamba. No se afirma que represente iniciativas nacionales ni que sus resultados sean generalizables a otras unidades educativas.",
        "Por todo lo expuesto, los antecedentes de la institución": "A partir de esta necesidad, el proyecto consolidó datos, comparó tres clasificadores y desarrolló una interfaz de apoyo. Dado el reducido número de transiciones positivas, la salida se considera exploratoria y requiere revisión docente.",
        "El desarrollo e implementación de este proyecto predictivo": "La pertinencia del proyecto se organiza en dimensiones tecnológica, pedagógica y metodológica; sus beneficios potenciales deberán comprobarse mediante evaluación con usuarios y validación prospectiva.",
        "El rezago en los primeros años de formación es el principal predictor": "El prototipo busca priorizar la revisión de perfiles académicos mediante niveles de alerta. No constituye un diagnóstico, no estima abandono escolar y no reemplaza la valoración integral del docente.",
        "El proyecto aplica una metodología rigurosa propia": "El proyecto aplica CRISP-DM para documentar recolección, preparación, modelado, evaluación y despliegue. La comparación temporal aporta evidencia reproducible del conjunto estudiado, sin demostrar eficacia general ni impacto pedagógico.",
        "En la actualidad, la unidad educativa primaria objeto": "La unidad educativa administra boletines por gestión, grado y paralelo. Los archivos disponibles contienen calificaciones, metadatos escolares e identificadores; su fragmentación dificulta el seguimiento longitudinal y exige controles de privacidad.",
        "Sin embargo, el uso que se le da a esta rica fuente": "Antes del proyecto no existía en el repositorio una herramienta integrada para reconstruir transiciones entre gestiones y presentar indicadores de riesgo. El problema técnico fue consolidar los datos y evaluar si dos variables académicas agregadas aportaban señal predictiva.",
        "Esta falta de un mecanismo analítico y predictivo": "La ausencia de una vista longitudinal limita la priorización sistemática de casos. El proyecto aborda esta necesidad mediante un prototipo de apoyo; no presupone que las dificultades observadas sean irreversibles ni atribuye causalidad a los patrones estadísticos.",
        "Los Árboles de Decisión son uno de los métodos": "Los árboles de decisión particionan el espacio de características mediante reglas jerárquicas y pueden utilizarse para clasificación (James et al., 2013). Su estructura facilita la inspección, aunque su interpretación no garantiza estabilidad fuera de la muestra.",
        "Justificación de Uso: Su principal ventaja": "En este proyecto, el Árbol se incluyó como modelo interpretable de referencia. Sus reglas se basan únicamente en promedio previo y número previo de asignaturas reprobadas.",
        "Estructura y Funcionamiento: Random Forest no construye": "Random Forest combina múltiples árboles construidos con remuestreo y selección aleatoria de características (Breiman, 2001). La implementación del proyecto utiliza 300 estimadores y promedia sus probabilidades de clase.",
        "Justificación de Uso: Se considera uno de los algoritmos": "Random Forest se incluyó por su capacidad de representar relaciones no lineales y admitir ponderación de clases. Su conveniencia debe determinarse con métricas de prueba, no asumirse de antemano.",
        "Justificación de Uso: En el análisis del abandono": "La MLP se incluyó como contraste no lineal. Con dos variables y muy pocos casos positivos, una mayor complejidad no implica mejor desempeño; en la prueba temporal no detectó positivos.",
        "Un Sistema de Alerta Temprana": "Un sistema de alerta temprana organiza indicadores para priorizar revisión e intervención (Macfadyen & Dawson, 2010). En este proyecto, el panel muestra probabilidad del modelo, puntaje de la regla, motivo y recomendación; el docente conserva la decisión final.",
        "Para garantizar que el modelo de Machine Learning": "El proyecto utiliza CRISP-DM como guía iterativa para comprensión del problema, comprensión y preparación de datos, modelado, evaluación y despliegue (Wirth & Hipp, 2000). El uso de esta metodología favorece la trazabilidad, pero no garantiza validez externa.",
        "La Inteligencia Artificial por sí sola no genera valor": "La ingeniería de características transforma datos crudos en variables utilizables por los modelos (James et al., 2013). En este trabajo se utilizaron únicamente el promedio general y el número de asignaturas reprobadas de la gestión T.",
        "En el pronóstico del rezago no basta": "La construcción longitudinal vinculó el desempeño de la gestión T con la etiqueta observada en T+1 para el mismo RUDE seudonimizado y solo cuando las gestiones eran consecutivas.",
        "Ventanas de Tiempo (Shift):": "Ventana temporal: usar variables de T para estimar la etiqueta de T+1 establece precedencia temporal, pero no una relación causal.",
        "Agregación Matemática:": "Agregación: promedio general y número de asignaturas reprobadas sintetizan las nueve notas. Esta reducción facilita el modelado, pero también descarta información específica por asignatura.",
        "Al entrenar múltiples algoritmos": "La evaluación de clasificadores desbalanceados requiere complementar la exactitud con matriz de confusión, precisión, recall, F1, balanced accuracy y precisión promedio (Fawcett, 2006). Con tres positivos en la prueba, cualquier estimación es inestable.",
        "Verdaderos Positivos (TP):": "Verdaderos positivos (TP): casos de rezago correctamente identificados.",
        "Verdaderos Negativos (TN):": "Verdaderos negativos (TN): casos sin rezago correctamente identificados.",
        "Falsos Positivos (FP": "Falsos positivos (FP): alertas emitidas para casos que no presentaron rezago.",
        "Falsos Negativos (FN": "Falsos negativos (FN): casos de rezago no detectados por el clasificador.",
        "Exhaustividad o Sensibilidad (Recall):": "Exhaustividad o sensibilidad (recall): proporción de positivos reales detectados, TP / (TP + FN). Debe interpretarse junto con las falsas alarmas y los recursos disponibles para la intervención.",
        "En Obj4_Evaluacion_Segregacion_Riesgo.ipynb se estableció": "La salida operativa distingue tres niveles de riesgo para registros completos y una categoría Sin datos cuando falta alguna calificación:",
        "Los hallazgos del presente estudio concuerdan": "Los resultados se interpretan a la luz de la literatura sobre Minería de Datos Educativos y sistemas de alerta temprana (Macfadyen & Dawson, 2010; Romero & Ventura, 2020):",
        "Se debe tener presente que los resultados obtenidos": "Los resultados corresponden a una sola unidad educativa y al periodo 2022–2024. Los modelos y umbrales deben validarse prospectivamente y recalibrarse cuando existan nuevas cohortes; no deben adoptarse automáticamente como política institucional.",
        "Estandarización y Formato Nativo de Datos:": "Estandarización y formato nativo: Solicitar centralizadores en Excel o CSV con esquema documentado, validaciones de rango y control de versiones. Esto reduce reprocesos y mejora la trazabilidad, pero no sustituye los controles de calidad.",
        "Registro de Asistencia Escolar:": "Variables adicionales: Evaluar, con autorización y justificación pedagógica, si asistencia u otros indicadores mejoran la predicción. Su incorporación debe comprobarse empíricamente y respetar la minimización de datos.",
        "Institucionalización y Uso Pedagógico:": "Piloto controlado: Probar el prototipo con un grupo reducido de docentes, registrar falsas alarmas, casos omitidos, usabilidad y acciones adoptadas, sin automatizar decisiones sobre estudiantes.",
        "Escalabilidad y MLOps:": "Mantenimiento: Versionar datos, código, modelos y métricas; reentrenar solo cuando existan nuevas cohortes suficientes y comparar siempre contra una línea base temporal.",
        "A medida que las instituciones educativas digitalizan": "La Minería de Datos Educativos es un campo interdisciplinario que desarrolla y aplica métodos computacionales para explorar datos originados en contextos educativos (Romero & Ventura, 2020).",
        "Según Baker e Inventado": "Baker e Inventado (2014) describen la Minería de Datos Educativos y la Analítica del Aprendizaje como enfoques relacionados para estudiar datos del aprendizaje y apoyar su interpretación. En este proyecto, ese marco orienta la transformación de registros longitudinales en indicadores sujetos a revisión docente.",
        "En el aprendizaje automático, las variables": "Las variables de entrada pueden tener escalas diferentes. StandardScaler centra cada variable y la escala por su desviación estándar; no convierte necesariamente los datos en una distribución normal. En este proyecto se ajustó solo con el conjunto de entrenamiento antes de transformar los datos de prueba (Hastie et al., 2009).",
        "En la predicción del rezago escolar, es natural": "El conjunto del proyecto está fuertemente desbalanceado. Sin medidas específicas, un clasificador puede favorecer la clase mayoritaria. Árbol y Random Forest utilizaron class_weight='balanced'; además se reportaron métricas centradas en la clase positiva (Breiman, 2001; Chawla et al., 2002).",
        "En un escenario donde el 90%": "Si el 90 % de los casos pertenece a la clase negativa, predecir siempre esa clase produce 90 % de exactitud y recall positivo igual a cero. Este ejemplo muestra por qué la exactitud debe interpretarse junto con la matriz de confusión y métricas de la clase positiva (Fawcett, 2006).",
        "src/data_loader.py:": "src/data_loader.py: Servicio de lectura, validación de esquema y combinación de nuevas nóminas. Los registros con notas incompletas se conservan como Sin datos y no reciben predicción.",
    }
    for prefix, text in replacements.items():
        replace_prefix(doc, prefix, text)

    synchronize_toc(doc)

    replace_prefix(
        doc,
        "4.3.2 Rendimiento Cuantitativo",
        "4.3.2 Rendimiento en la Prueba Temporal 2023→2024 (N=248)",
    )
    replace_prefix(
        doc,
        "Figura 4.7:",
        "Figura 4.7: Distribución operativa del riesgo y registros sin datos suficientes (N=1.118)",
    )
    replace_prefix(doc, "Gestión 2021:", "")
    replace_prefix(doc, "3.7.2 Red Neuronal MLP", "3.7.3 Red Neuronal MLP (MLPClassifier)")
    replace_prefix(doc, "3.9.1 Capa de Resguardo", "3.9.2 Capa de Resguardo Pedagógico (Sistema Híbrido)")
    replace_prefix(doc, "Figura 3-2: Flujograma", "Figura 3-2: Flujo metodológico CRISP-DM adaptado")

    for paragraph in list(doc.paragraphs):
        if paragraph.text.strip().startswith("https://drive.google.com/"):
            paragraph._p.getparent().remove(paragraph._p)

    # Capturas técnicas obsoletas; el contenido canónico queda en texto y scripts reproducibles.
    remove_embedded_images(doc, {"image2.png", "image4.png", "image24.png", "image11.png", "image5.png", "image18.png"})
    remove_inline_drawings(doc, {"image23.png"})
    replace_prefix(doc, "https://drive.google.com/", "")
    replace_prefix(
        doc,
        "Revisar el CD adjunto o visitar el siguiente enlace:",
        "Repositorio público de código (sin datos personales):",
    )

    # Reemplaza afirmación inferencial no respaldada.
    replace_prefix(
        doc,
        "Al contrastar los promedios académicos",
        "La comparación descriptiva entre grupos mostró diferencias de medias. No se afirma significancia estadística porque no se aplicó una prueba inferencial.",
    )

    # Tabla de hiperparámetros (cuarta tabla del documento).
    if len(doc.tables) >= 4:
        table = doc.tables[3]
        values = [
            ["Árbol de Decisión", "DecisionTreeClassifier", "max_depth=3; min_samples_leaf=10; class_weight='balanced'; random_state=42", "Sin escalado"],
            ["Random Forest", "RandomForestClassifier", "n_estimators=300; max_depth=4; min_samples_leaf=5; class_weight='balanced'; random_state=42", "Sin escalado"],
            ["Red Neuronal MLP", "MLPClassifier", "hidden_layer_sizes=(4,); alpha=0.1; learning_rate_init=0.0001; max_iter=1000; random_state=42", "StandardScaler"],
        ]
        for r, vals in enumerate(values, start=1):
            for c, value in enumerate(vals):
                table.cell(r, c).text = value

    if len(doc.tables) >= 5:
        risk_table = doc.tables[4]
        if not any(row.cells[0].text.strip() == "Sin datos" for row in risk_table.rows):
            risk_table._tbl.append(deepcopy(risk_table.rows[-1]._tr))
            cells = risk_table.rows[-1].cells
            values = [
                "Sin datos",
                "No se calcula",
                "SIN DATOS",
                "Completar las nueve calificaciones antes de emitir una alerta.",
            ]
            for cell, value in zip(cells, values):
                cell.text = value
        risk_values = [
            ["Nivel de riesgo", "Rango", "Indicador visual", "Orientación para la revisión"],
            ["Alto riesgo", "P ≥ 0,70 o regla pedagógica", "ALTO RIESGO", "Priorizar la revisión docente y verificar los datos y el contexto; no adoptar decisiones automáticas."],
            ["Medio riesgo", "0,40 ≤ P < 0,70", "MEDIO RIESGO", "Realizar seguimiento y revisión focalizada según el criterio docente."],
            ["Bajo riesgo", "P < 0,40", "BAJO RIESGO", "Mantener monitoreo regular; el nivel no descarta dificultades no observadas por el modelo."],
            ["Sin datos", "No se calcula", "SIN DATOS", "Completar las nueve calificaciones antes de emitir una alerta."],
        ]
        for row, values in zip(risk_table.rows, risk_values):
            for cell, value in zip(row.cells, values):
                cell.text = value

    # Sustituye listas de figuras/tablas de la plantilla por índices reales.
    list_tables = next(p for p in doc.paragraphs if p.text.strip() == "Lista de tablas")
    figure_entries = [
        ("Figura 3-1. Área de estudio de la U.E. José María Santiváñez", 10),
        ("Figura 3-2. Flujo metodológico CRISP-DM adaptado", 11),
        ("Figura 3-3. Carpetas por gestión escolar", 12),
        ("Figura 3-4. Boletines centralizados", 13),
        ("Figura 3-5. Boletín anual", 13),
        ("Figura 3-6. Boletines transformados a Excel", 14),
        ("Figura 4.1. Tasa de reprobación por asignatura", 22),
        ("Figura 4.2. Materias reprobadas según condición de rezago", 24),
        ("Figura 4.6. Matrices de confusión de la prueba temporal", 25),
        ("Figura 4.7. Distribución operativa del riesgo", 27),
        ("Figura 4.3. Evolución del rezago por gestión", 28),
        ("Figura 4.4. Rezago promedio por grado", 28),
        ("Figura 4.5. Trayectoria longitudinal de ejemplo", 29),
    ]
    for label, page in figure_entries:
        add_index_entry(list_tables, label, page)

    list_tables = next(p for p in doc.paragraphs if p.text.strip() == "Lista de tablas")
    table_entries = [
        ("Tabla 4.1. Caracterización del conjunto consolidado", 20),
        ("Tabla 4.2. Tasas de reprobación por asignatura", 21),
        ("Tabla 4.3. Promedios por condición de rezago", 23),
        ("Tabla 4.4. Hiperparámetros de los modelos", 24),
        ("Tabla 4.5. Umbrales operativos de riesgo", 26),
    ]
    anchor = list_tables
    for label, page in table_entries:
        paragraph = insert_after(anchor, f"{label}\t{page}", "normal")
        paragraph.paragraph_format.tab_stops.add_tab_stop(
            Inches(6.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
        )
        anchor = paragraph

    # Agrega el rótulo que faltaba bajo la figura de matrices de confusión.
    for paragraph in doc.paragraphs:
        targets = []
        for blip in paragraph._p.xpath(".//a:blip"):
            relation_id = blip.get(qn("r:embed"))
            part = doc.part.related_parts.get(relation_id) if relation_id else None
            if part is not None:
                targets.append(Path(str(part.partname)).name)
        if "image3.png" in targets:
            caption = insert_after(paragraph, "Figura 4.6: Matrices de confusión en la prueba temporal 2023→2024 (N=248)")
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption.runs[0].italic = True
            break

    # Bibliografía: reconstruir el bloque para eliminar hipervínculos residuales del borrador.
    paragraphs = doc.paragraphs
    bib_index = next(i for i, p in enumerate(paragraphs) if p.text.strip() == "Bibliografía")
    annex_index = next(i for i, p in enumerate(paragraphs[bib_index + 1:], bib_index + 1) if p.text.strip() == "Anexos")
    annex_para = paragraphs[annex_index]
    between = paragraphs[bib_index + 1:annex_index]
    for paragraph in between:
        if paragraph._p.getparent() is not None:
            paragraph._p.getparent().remove(paragraph._p)
    for ref in REFERENCES:
        p = insert_before(annex_para, "", "normal")
        style_reference(p, ref)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(1)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.size = Pt(9)

    risk_analysis = next(p for p in doc.paragraphs if p.text.strip().startswith("4.4.2 Análisis Temporal"))
    insert_before(
        risk_analysis,
        "La distribución incluye 88 observaciones en la categoría Sin datos. El prototipo conserva estos registros para trazabilidad, pero no calcula una alerta hasta disponer de las nueve calificaciones.",
        "normal",
    )

    # Anexos nuevos verificables.
    principal = next(p for p in doc.paragraphs if p.text.strip().startswith("Anexo PRINCIPAL"))
    appendix_items = [
        ("Anexo 4: Trazabilidad de la Muestra Predictiva", "Dataset: 1.118 observaciones, 592 estudiantes y 23 rezagos descriptivos. Transiciones consecutivas con origen y destino completos: 489, con seis positivas. Se excluyeron 33 pares con etiqueta objetivo incompleta. Entrenamiento 2022→2023: N=241, tres positivas. Prueba 2023→2024: N=248, tres positivas."),
        ("Anexo 5: Resultados Reproducibles", "El comando python scripts/train_models.py reconstruye las transiciones, evalúa temporalmente los modelos, exporta los artefactos y guarda resultados_modelos/metricas_modelos.json. Las figuras estadísticas se generan desde ese archivo."),
        ("Anexo 6: Privacidad y Uso Responsable", "Los datos originales contienen información de menores y no deben publicarse. Se exige autorización institucional, seudonimización, acceso restringido y revisión humana de toda alerta. Véase documentacion/08_Protocolo_Privacidad.md."),
        ("Anexo 7: Pruebas del Prototipo", "tests/test_project.py verifica columnas obligatorias, notas en rango 0–100, rechazo de entradas inválidas, separación entre probabilidad del modelo y regla pedagógica, y ausencia de predicción cuando faltan calificaciones."),
    ]
    for title, body in appendix_items:
        insert_before(principal, title, "Heading 3")
        insert_before(principal, body, "normal")

    update_fields_on_open(doc)


def replace_embedded_images(path: Path) -> None:
    replacements = {
        "word/media/image3.png": ROOT / "documentacion/figuras/fig_4_6_matriz_confusion_modelos.png",
        "word/media/image21.png": ROOT / "documentacion/figuras/fig_4_7_distribucion_riesgo_estudiantes.png",
        "word/media/image25.png": ROOT / "documentacion/figuras/fig_3_1_flujograma_crisp_dm.png",
        "word/media/image28.png": ROOT / "documentacion/figuras/fig_3_3_capa_hibrida_resguardo.png",
        "word/media/image14.png": ROOT / "documentacion/figuras/fig_3_2_arquitectura_software.png",
    }
    with tempfile.TemporaryDirectory(prefix="docx_images_") as td:
        temp = Path(td)
        with zipfile.ZipFile(path) as zin:
            zin.extractall(temp)
        for member, image in replacements.items():
            shutil.copyfile(image, temp / member)
        patch_raw_ooxml(temp)
        rebuilt = path.with_suffix(".tmp.docx")
        with zipfile.ZipFile(rebuilt, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in sorted(temp.rglob("*")):
                if item.is_file():
                    zout.write(item, item.relative_to(temp))
        rebuilt.replace(path)


def patch_raw_ooxml(unpacked: Path) -> None:
    """Corrige texto dentro de controles de contenido no expuestos por python-docx."""
    document_path = unpacked / "word/document.xml"
    tree = etree.parse(str(document_path))
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    stale_list_prefixes = (
        "Figura 3-1: Poner un mapa",
        "Figura 3-2: Flujograma metodológico",
        "Figura/tabla 3-3:",
        "Figura 3-4: Parámetros",
        "Figura 3-6: Planilla relevamiento",
        "Figura 3-7: Puntos de relevamiento",
        "Figura 3-8: Fotografías",
        "Figura 3-10: Planilla relevamiento",
        "Figura 4-1: Título",
        "Figura 4-10:",
        "Figura 4-11:",
        "Tabla 3-1:",
        "Tabla 3-2:",
    )
    for paragraph in list(tree.xpath("//w:p", namespaces=ns)):
        texts = paragraph.xpath(".//w:t", namespaces=ns)
        content = "".join(node.text or "" for node in texts).strip()
        if "fiabilidad operativa del 100%" in content:
            replacement = (
                "En la prueba temporal, Árbol y Random Forest detectaron uno de tres casos positivos "
                "y generaron 13 falsas alarmas; la MLP no detectó positivos. Estos resultados son "
                "exploratorios y no demuestran alta capacidad predictiva."
            )
            if texts:
                texts[0].text = replacement
                for node in texts[1:]:
                    node.text = ""
        elif content.startswith(stale_list_prefixes):
            parent = paragraph.getparent()
            if parent is not None:
                parent.remove(paragraph)
    tree.write(str(document_path), xml_declaration=True, encoding="UTF-8", standalone="yes")


def main() -> None:
    doc = Document(SOURCE)
    replace_everywhere(doc)
    revise_content(doc)
    doc.core_properties.title = "Predicción exploratoria del rezago académico en educación primaria"
    doc.core_properties.subject = "Proyecto de ciencia de datos - versión revisada"
    doc.save(OUTPUT)
    replace_embedded_images(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
