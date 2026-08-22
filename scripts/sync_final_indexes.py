"""Sincroniza índices y numeración tras la revisión visual final."""

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "documentacion/Proyecto_Revisado_Academicamente_Final.docx"

TOC_PAGES = {
    "3.10.": 19,
    "3.10.1.": 19,
    "3.10.2.": 20,
    "3.10.3.": 20,
    "3.10.4.": 21,
    "3.10.5.": 22,
    "4.": 24,
    "4.1.": 24,
    "4.1.1.": 24,
    "4.1.2.": 25,
    "4.2.": 25,
    "4.2.1.": 25,
    "4.2.2.": 27,
    "4.3.": 28,
    "4.3.1.": 28,
    "4.3.2.": 29,
    "4.4.": 30,
    "4.4.1.": 30,
    "4.4.2.": 31,
    "4.5.": 33,
    "4.5.1.": 35,
    "4.6.": 35,
    "5.": 37,
    "6.": 38,
    "Bibliografía": 39,
    "Anexos": 41,
}

FIGURES = [
    ("Figura 3.1. Área de estudio de la U.E. José María Santiváñez", 10),
    ("Figura 3.2. Flujo metodológico CRISP-DM adaptado", 11),
    ("Figura 3.3. Carpetas por gestión escolar", 12),
    ("Figura 3.4. Boletines centralizados", 13),
    ("Figura 3.5. Boletín anual", 13),
    ("Figura 3.6. Boletines transformados a Excel", 14),
    ("Figura 3.7. Arquitectura modular y flujo de ejecución del prototipo", 17),
    ("Figura 3.8. Lógica de inferencia y resguardo pedagógico del sistema híbrido", 18),
    ("Figura 3.9. Panel principal del prototipo con filtros e indicadores agregados", 22),
    ("Figura 3.10. Simulador libre y separación entre probabilidad y puntaje operativo", 23),
    ("Figura 4.1. Tasa de reprobación por asignatura", 26),
    ("Figura 4.2. Materias reprobadas según condición de rezago", 28),
    ("Figura 4.3. Matrices de confusión de la prueba temporal", 29),
    ("Figura 4.4. Distribución operativa del riesgo", 31),
    ("Figura 4.5. Evolución del rezago por gestión", 32),
    ("Figura 4.6. Rezago promedio por grado", 32),
    ("Figura 4.7. Trayectoria longitudinal de ejemplo", 33),
    ("Figura 4.8. Vista de monitoreo del curso y panel de alertas tempranas", 34),
    ("Figura 4.9. Simulador libre para comprobar escenarios y reglas operativas", 35),
]

TABLES = [
    ("Tabla 3.1. Matriz de trazabilidad entre objetivos, procesos y evidencias", 19),
    ("Tabla 3.2. Reglas de validación y tratamiento de datos académicos", 20),
    ("Tabla 3.3. Construcción y partición de la muestra predictiva longitudinal", 21),
    ("Tabla 4.1. Caracterización del conjunto consolidado", 24),
    ("Tabla 4.2. Tasas de reprobación por asignatura", 25),
    ("Tabla 4.3. Promedios por condición de rezago", 27),
    ("Tabla 4.4. Hiperparámetros de los modelos", 29),
    ("Tabla 4.5. Umbrales operativos de riesgo", 30),
]


def remove_paragraph(paragraph) -> None:
    parent = paragraph._p.getparent()
    if parent is not None:
        parent.remove(paragraph._p)


def restyle(paragraph, size=10.5) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(1)
    paragraph.paragraph_format.line_spacing = 1.0
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.size = Pt(size)


def replace_caption(doc, old_prefix: str, new_text: str) -> None:
    paragraphs = list(doc.paragraphs)
    for paragraph in paragraphs:
        if not paragraph.text.strip().startswith(old_prefix):
            continue
        idx = next(i for i, item in enumerate(paragraphs) if item._p is paragraph._p)
        if any(item._p.xpath(".//a:blip") for item in paragraphs[max(0, idx - 4):idx]):
            paragraph.text = new_text
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.name = "Arial"
                run.font.size = Pt(10)
                run.italic = True
            return
    raise ValueError(f"No se encontró el pie corporal: {old_prefix}")


def rebuild_list(doc, start_title: str, end_title: str, entries) -> None:
    paragraphs = list(doc.paragraphs)
    start = next(i for i, p in enumerate(paragraphs) if p.text.strip() == start_title)
    end = next(i for i, p in enumerate(paragraphs) if p.text.strip() == end_title)
    existing = [p for p in paragraphs[start + 1:end] if p.text.strip()]
    template_ppr = deepcopy(existing[0]._p.pPr) if existing and existing[0]._p.pPr is not None else None
    between = paragraphs[start + 1:end]
    section_break = next(
        (
            paragraph for paragraph in between
            if paragraph._p.find("./w:pPr/w:sectPr", paragraph._p.nsmap) is not None
        ),
        None,
    )
    anchor = section_break if section_break is not None else paragraphs[end]
    for paragraph in between:
        if section_break is not None and paragraph._p is section_break._p:
            continue
        remove_paragraph(paragraph)
    for title, page in entries:
        paragraph = doc.add_paragraph()
        if template_ppr is not None:
            if paragraph._p.pPr is not None:
                paragraph._p.remove(paragraph._p.pPr)
            paragraph._p.insert(0, deepcopy(template_ppr))
        paragraph.text = f"{title}\t{page}"
        paragraph.paragraph_format.tab_stops.add_tab_stop(
            Inches(6.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
        )
        restyle(paragraph)
        anchor._p.addprevious(paragraph._p)


def update_toc(doc) -> None:
    paragraphs = list(doc.paragraphs)
    fig_heading = next(i for i, p in enumerate(paragraphs) if p.text.strip() == "Lista de figuras")
    for paragraph in paragraphs[:fig_heading]:
        text = paragraph.text.strip()
        if not text or "\t" not in text:
            continue
        key = text.split("\t", 1)[0].strip()
        if key not in TOC_PAGES:
            continue
        title = "\t".join(text.split("\t")[:-1])
        paragraph.text = f"{title}\t{TOC_PAGES[key]}"
        restyle(paragraph, size=11)


def main() -> None:
    doc = Document(DOCX)

    # Renumeración por orden real de aparición en el capítulo 4.
    replace_caption(doc, "Figura 4.6. Matrices", "Figura 4.3. Matrices de confusión en la prueba temporal 2023→2024 (N=248)")
    replace_caption(doc, "Figura 4.7. Distribución", "Figura 4.4. Distribución operativa del riesgo y registros sin datos suficientes (N=1.118)")
    replace_caption(doc, "Figura 4.3. Evolución", "Figura 4.5. Evolución de la proporción de rezago académico (2022–2024)")
    replace_caption(doc, "Figura 4.4. Proporción", "Figura 4.6. Proporción de rezago promedio por grado escolar")
    replace_caption(doc, "Figura 4.5. Evolución Académica", "Figura 4.7. Evolución académica longitudinal de un estudiante de ejemplo")

    update_toc(doc)
    rebuild_list(doc, "Lista de figuras", "Lista de tablas", FIGURES)
    rebuild_list(doc, "Lista de tablas", "Introducción", TABLES)
    doc.save(DOCX)
    print(DOCX)


if __name__ == "__main__":
    main()
