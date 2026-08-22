"""Amplía el documento académico con evidencia metodológica reproducible.

El archivo de entrada se conserva intacto. La salida añade una sección de
trazabilidad por objetivo, evidencia visual anonimizada, pies para todas las
tablas existentes y un anexo de reproducibilidad.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "documentacion/Proyecto_Revisado_Academicamente.docx"
OUTPUT = ROOT / "documentacion/Proyecto_Revisado_Academicamente_Ampliado.docx"
FIG_DIR = ROOT / "documentacion/figuras"

BLUE = "2455D6"
LIGHT_BLUE = "EAF2FF"
GRID = "B7C7DB"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=70, start=80, bottom=70, end=80) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table, color: str = GRID, size: str = "6") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def keep_together(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    keep_lines = OxmlElement("w:keepLines")
    keep_lines.set(qn("w:val"), "true")
    p_pr.append(keep_lines)


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.15) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def style_body(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_spacing(paragraph, after=6, line=1.15)
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.size = Pt(11)


def insert_paragraph_before(anchor, text="", style=None, *, bold_prefix=None):
    p = anchor.insert_paragraph_before(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        p.add_run(bold_prefix).bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    style_body(p)
    if style and style.startswith("Heading"):
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.keep_with_next = True
    return p


def insert_page_break_before(anchor) -> None:
    p = anchor.insert_paragraph_before()
    p.add_run().add_break(WD_BREAK.PAGE)


def insert_picture_before(anchor, path: Path, width=Inches(6.35)):
    p = anchor.part.document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(path), width=width)
    anchor._p.addprevious(p._p)
    return p


def insert_caption_before(anchor, label: str, title: str, source: str):
    cap = anchor.insert_paragraph_before()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(2)
    cap.paragraph_format.space_after = Pt(1)
    cap.paragraph_format.keep_with_next = True
    run = cap.add_run(f"{label}. {title}")
    run.font.name = "Arial"
    run.font.size = Pt(10)
    run.italic = True
    keep_together(cap)

    src = anchor.insert_paragraph_before()
    src.alignment = WD_ALIGN_PARAGRAPH.CENTER
    src.paragraph_format.space_before = Pt(0)
    src.paragraph_format.space_after = Pt(7)
    run = src.add_run(f"Fuente: {source}")
    run.font.name = "Arial"
    run.font.size = Pt(9)
    keep_together(src)
    return cap, src


def build_table_before(anchor, headers, rows, widths=None):
    doc = anchor.part.document
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    if widths:
        for i, width in enumerate(widths):
            table.columns[i].width = Inches(width)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = text
        set_cell_shading(cell, BLUE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            for run in p.runs:
                run.font.name = "Arial"
                run.font.size = Pt(9)
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
    for ridx, row_data in enumerate(rows, start=1):
        row = table.add_row()
        for i, value in enumerate(row_data):
            cell = row.cells[i]
            cell.text = str(value)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if ridx % 2 == 0:
                set_cell_shading(cell, LIGHT_BLUE)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                for run in p.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(8.5)
    anchor._p.addprevious(table._tbl)
    return table


def add_table_caption_after(table, label: str, title: str, source: str) -> None:
    doc = table._parent
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(2)
    cap.paragraph_format.space_after = Pt(1)
    r = cap.add_run(f"{label}. {title}")
    r.font.name = "Arial"
    r.font.size = Pt(10)
    r.italic = True
    table._tbl.addnext(cap._p)

    src = doc.add_paragraph()
    src.alignment = WD_ALIGN_PARAGRAPH.CENTER
    src.paragraph_format.space_before = Pt(0)
    src.paragraph_format.space_after = Pt(7)
    r = src.add_run(f"Fuente: {source}")
    r.font.name = "Arial"
    r.font.size = Pt(9)
    cap._p.addnext(src._p)


def find_paragraph(doc, exact: str):
    return next(p for p in doc.paragraphs if p.text.strip() == exact)


def insert_list_entries(anchor, entries) -> None:
    for entry in entries:
        p = anchor.insert_paragraph_before(entry)
        p.style = anchor.style
        p.paragraph_format.left_indent = anchor.paragraph_format.left_indent
        p.paragraph_format.first_line_indent = anchor.paragraph_format.first_line_indent
        for run in p.runs:
            run.font.name = "Arial"
            run.font.size = Pt(11)


def normalize_existing_figure_captions(doc) -> None:
    pattern = re.compile(r"^Figura\s+(\d+)[.-](\d+)\s*:\s*(.+)$")
    for p in doc.paragraphs:
        match = pattern.match(p.text.strip())
        if not match:
            continue
        p.text = f"Figura {match.group(1)}.{match.group(2)}. {match.group(3).strip()}"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.name = "Arial"
            run.font.size = Pt(10)
            run.italic = True

    for p in doc.paragraphs:
        text = p.text.strip()
        if text == "Fuente: Elaboración propia 2026":
            p.text = "Fuente: elaboración propia (2026)."
        elif text == "Fuente: Google Maps 2026":
            p.text = "Fuente: Google Maps (2026)."
        if p.text.startswith("Fuente:"):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = "Arial"
                run.font.size = Pt(9)


def add_methodology_section(doc) -> None:
    anchor = find_paragraph(doc, "Análisis de Resultados y Discusión")
    insert_paragraph_before(anchor, "3.10 Trazabilidad metodológica y evidencia por objetivo específico", "Heading 2")
    insert_paragraph_before(
        anchor,
        "Para fortalecer la reproducibilidad del capítulo, esta sección relaciona cada objetivo específico con su entrada, procedimiento, control y evidencia verificable. Las imágenes de código muestran fragmentos de la implementación canónica y se acompañan de salidas agregadas; no exponen nombres ni identificadores estudiantiles. La Tabla 3.1 resume la correspondencia entre los objetivos y los productos generados.",
    )

    table = build_table_before(
        anchor,
        ["Objetivo", "Proceso ejecutado", "Evidencia verificable", "Producto"],
        [
            ("OE1", "Ingesta, validación, normalización y consolidación de boletines.", "Controles de esquema, rango, duplicidad y completitud.", "Dataset consolidado y registros Sin datos."),
            ("OE2", "Análisis descriptivo por asignatura y condición de rezago.", "Tasas, promedios y gráficos agregados.", "Patrones descriptivos de la muestra."),
            ("OE3", "Construcción T→T+1, partición temporal y entrenamiento de tres clasificadores.", "Muestra predictiva, hiperparámetros y métricas de prueba.", "Modelos y archivo de métricas reproducibles."),
            ("OE4", "Evaluación y categorización operativa del riesgo.", "Matriz de confusión, distribución de categorías y motivo de alerta.", "Priorización para revisión docente."),
            ("OE5", "Integración de datos, modelos y controles en Streamlit.", "Capturas del panel y del simulador libre.", "Prototipo funcional de apoyo."),
        ],
        widths=[0.75, 2.0, 1.9, 1.7],
    )
    add_table_caption_after(
        table,
        "Tabla 3.1",
        "Matriz de trazabilidad entre objetivos, procesos y evidencias",
        "elaboración propia a partir de los artefactos del proyecto (2026).",
    )

    insert_paragraph_before(anchor, "3.10.1 Evidencia del objetivo específico 1: calidad y consolidación", "Heading 3")
    insert_paragraph_before(
        anchor,
        "La entrada estuvo constituida por 36 boletines organizados por gestión, grado y paralelo. Después de la conversión a planillas, el cargador verifica columnas obligatorias, convierte las nueve calificaciones a formato numérico, rechaza valores fuera del intervalo 0–100 y controla la unicidad de la combinación estudiante–gestión. Los faltantes se conservan explícitamente: no se imputaron notas ni se asignó una categoría de riesgo a registros incompletos.",
    )
    insert_page_break_before(anchor)
    table = build_table_before(
        anchor,
        ["Control", "Criterio implementado", "Respuesta ante incumplimiento", "Evidencia"],
        [
            ("Esquema", "Presencia de gestión, curso, paralelo, identificador y nueve asignaturas.", "Se detiene la carga y se informa qué columnas faltan.", "src/data_loader.py"),
            ("Tipo", "Calificaciones convertibles a valores numéricos.", "El dato no convertible queda marcado como faltante.", "validate_dataframe()"),
            ("Rango", "Cada nota debe estar entre 0 y 100.", "Se rechaza el conjunto cargado.", "Prueba automatizada"),
            ("Unicidad", "Una observación por estudiante y gestión.", "Se reporta la duplicidad para corrección.", "Clave estudiante–gestión"),
            ("Completitud", "Nueve notas presentes para calcular promedio y reprobadas.", "Se conserva el registro como Sin datos y no se predice.", "1.030 completos; 88 incompletos"),
        ],
        widths=[0.85, 2.15, 2.15, 1.2],
    )
    add_table_caption_after(
        table,
        "Tabla 3.2",
        "Reglas de validación y tratamiento de datos académicos",
        "elaboración propia con base en src/data_loader.py y las pruebas automatizadas (2026).",
    )
    insert_picture_before(anchor, FIG_DIR / "fig_3_7_obj1_validacion_consolidacion.png")
    insert_caption_before(
        anchor,
        "Figura 3.7",
        "Validación, construcción de variables y resultado agregado del objetivo 1",
        "elaboración propia mediante Obj1_Recoleccion_Limpieza.ipynb y src/data_loader.py (2026).",
    )
    insert_paragraph_before(
        anchor,
        "El resultado verificable fue un conjunto de 1.118 observaciones estudiante-año de 592 estudiantes. De ellas, 1.030 contienen las nueve calificaciones requeridas y 88 permanecen identificadas como incompletas. Esta distinción evita transformar la ausencia de información en una supuesta aprobación o reprobación.",
    )

    insert_paragraph_before(anchor, "3.10.2 Evidencia del objetivo específico 2: análisis de patrones", "Heading 3")
    insert_paragraph_before(
        anchor,
        "El análisis exploratorio calculó, para cada asignatura, la proporción de observaciones con nota inferior a 51 y comparó promedios por condición de rezago. En el cálculo porcentual, las notas faltantes no se contabilizan como reprobadas; por ello las tasas describen el dataset disponible y deben leerse junto con el control de completitud. No se realizaron pruebas causales ni se atribuyeron los resultados a características personales.",
    )
    insert_picture_before(anchor, FIG_DIR / "fig_3_8_obj2_analisis_patrones.png")
    insert_caption_before(
        anchor,
        "Figura 3.8",
        "Código reproducible y tasas descriptivas de reprobación por asignatura",
        "elaboración propia con datos agregados de Obj2a_Analisis_Patrones.ipynb (2026).",
    )
    insert_paragraph_before(
        anchor,
        "La salida identifica a Comunicación y Lenguajes (2,06 %) y Matemática (1,70 %) como las asignaturas con mayor tasa observada. El hallazgo cumple una función de focalización descriptiva: señala dónde conviene revisar evidencias pedagógicas, pero no prueba que una asignatura cause el rezago.",
    )

    insert_paragraph_before(anchor, "3.10.3 Evidencia del objetivo específico 3: modelado temporal", "Heading 3")
    insert_paragraph_before(
        anchor,
        "La unidad de análisis predictiva no fue la fila aislada, sino la transición consecutiva del mismo estudiante entre T y T+1. Las variables de entrada fueron el promedio general y el número de asignaturas reprobadas en T; la etiqueta fue la presencia de rezago en T+1. Se exigieron datos completos en ambos extremos de la transición y se separó el entrenamiento 2022→2023 de la prueba 2023→2024 para evitar utilizar información futura al evaluar.",
    )
    table = build_table_before(
        anchor,
        ["Componente", "Periodo", "N", "Positivos", "Uso"],
        [
            ("Dataset descriptivo", "2022–2024", "1.118", "23 observaciones con rezago", "Caracterización"),
            ("Transiciones válidas", "T→T+1", "489", "6", "Modelado"),
            ("Entrenamiento", "2022→2023", "241", "3", "Ajuste de modelos"),
            ("Prueba temporal", "2023→2024", "248", "3", "Evaluación principal"),
            ("Pares excluidos", "T→T+1", "33", "No aplicable", "Destino incompleto"),
        ],
        widths=[1.4, 1.15, 0.75, 1.4, 1.65],
    )
    add_table_caption_after(
        table,
        "Tabla 3.3",
        "Construcción y partición de la muestra predictiva longitudinal",
        "elaboración propia a partir de scripts/train_models.py y resultados_modelos/metricas_modelos.json (2026).",
    )
    insert_picture_before(anchor, FIG_DIR / "fig_3_9_obj3_entrenamiento_temporal.png")
    insert_caption_before(
        anchor,
        "Figura 3.9",
        "Entrenamiento reproducible y evaluación temporal de los tres clasificadores",
        "elaboración propia a partir de Obj3a, Obj3b y scripts/train_models.py (2026).",
    )
    insert_paragraph_before(
        anchor,
        "En prueba, Árbol y Random Forest detectaron uno de los tres positivos, con precisión 0,0714, recall 0,3333, F1 0,1176 y exactitud balanceada 0,6401. La MLP no detectó positivos. Debido a que la prueba contiene solamente tres eventos de rezago, las métricas se reportan como exploratorias y no permiten afirmar generalización ni superioridad concluyente.",
    )

    insert_paragraph_before(anchor, "3.10.4 Evidencia del objetivo específico 4: evaluación y segregación", "Heading 3")
    insert_paragraph_before(
        anchor,
        "La aplicación conserva dos salidas diferentes: la probabilidad estimada por el modelo y un puntaje operativo utilizado para ordenar la revisión. Cuando el promedio es inferior a 51 o existen dos o más asignaturas reprobadas, una regla pedagógica eleva el puntaje de revisión; esta modificación no se presenta como probabilidad calibrada. Cada registro conserva el motivo de la alerta y los incompletos se clasifican como Sin datos.",
    )
    insert_page_break_before(anchor)
    insert_picture_before(anchor, FIG_DIR / "fig_3_10_obj4_segregacion_riesgo.png")
    insert_caption_before(
        anchor,
        "Figura 3.10",
        "Trazabilidad de la segregación operativa y distribución de categorías",
        "elaboración propia mediante Obj4_Evaluacion_Segregacion_Riesgo.ipynb y src/predictor.py (2026).",
    )
    insert_paragraph_before(
        anchor,
        "La ejecución sobre las 1.118 observaciones produjo 881 registros en Bajo riesgo, 116 en Medio riesgo, 33 en Alto riesgo y 88 Sin datos. Esta distribución describe la lógica operativa actual; no equivale a prevalencia futura de rezago ni valida la eficacia de una intervención.",
    )

    insert_paragraph_before(anchor, "3.10.5 Evidencia del objetivo específico 5: prototipo docente", "Heading 3")
    insert_paragraph_before(
        anchor,
        "El prototipo integra filtros por gestión, grado y paralelo; tarjetas agregadas; una tabla de seguimiento; simulación por estudiante y simulación libre. La interfaz distingue el resultado estadístico del criterio operativo y muestra Sin datos cuando no existe información suficiente. Las capturas siguientes se obtuvieron de la aplicación ejecutada localmente y fueron anonimizadas para no exponer datos de estudiantes.",
    )
    insert_picture_before(anchor, FIG_DIR / "fig_3_10_app_panel_principal.png")
    insert_caption_before(
        anchor,
        "Figura 3.11",
        "Panel principal del prototipo con filtros e indicadores agregados",
        "captura anonimizada de la aplicación Streamlit desarrollada en el proyecto (2026).",
    )
    insert_paragraph_before(
        anchor,
        "El panel principal permite delimitar la cohorte analizada y consultar cantidades por nivel de riesgo sin convertir la alerta en una decisión automática. La recomendación mostrada es una pauta para revisión profesional y debe contrastarse con información pedagógica adicional.",
    )
    insert_picture_before(anchor, FIG_DIR / "fig_3_11_app_simulador_libre.png")
    insert_caption_before(
        anchor,
        "Figura 3.12",
        "Simulador libre y separación entre probabilidad del modelo y puntaje operativo",
        "captura anonimizada de la aplicación Streamlit desarrollada en el proyecto (2026).",
    )
    insert_paragraph_before(
        anchor,
        "El simulador libre permite comprobar escenarios de entrada sin modificar el dataset. Su finalidad es didáctica y de verificación de reglas; no debe emplearse para etiquetar estudiantes reales fuera del flujo institucional de validación, autorización y revisión humana.",
    )
    insert_page_break_before(anchor)


def add_existing_table_captions(doc) -> None:
    specs = [
        ("Tabla 4.1", "Caracterización del conjunto consolidado", "elaboración propia a partir de primaria_dataset.csv (2026)."),
        ("Tabla 4.2", "Tasas descriptivas de reprobación por asignatura", "elaboración propia a partir del dataset consolidado (2026)."),
        ("Tabla 4.3", "Promedios por condición de rezago", "elaboración propia a partir del análisis exploratorio (2026)."),
        ("Tabla 4.4", "Hiperparámetros de los modelos evaluados", "elaboración propia con base en scripts/train_models.py (2026)."),
        ("Tabla 4.5", "Umbrales de clasificación operativa del riesgo", "elaboración propia con base en config.py y src/predictor.py (2026)."),
    ]
    for table, (label, title, source) in zip(list(doc.tables)[:5], specs):
        add_table_caption_after(table, label, title, source)


def add_existing_result_figure_sources(doc) -> None:
    specs = {
        "Figura 4.1.": "elaboración propia a partir del dataset consolidado (2026).",
        "Figura 4.2.": "elaboración propia a partir del análisis exploratorio (2026).",
        "Figura 4.3.": "elaboración propia a partir del dataset consolidado (2026).",
        "Figura 4.4.": "elaboración propia a partir del dataset consolidado (2026).",
        "Figura 4.5.": "elaboración propia a partir de la trayectoria longitudinal seudonimizada (2026).",
        "Figura 4.6.": "elaboración propia a partir de resultados_modelos/metricas_modelos.json (2026).",
        "Figura 4.7.": "elaboración propia a partir del dataset consolidado y src/predictor.py (2026).",
    }
    in_body = False
    for p in list(doc.paragraphs):
        if p.text.strip() == "Ingeniería del proyecto" and p.style.name == "Heading 1":
            in_body = True
        if not in_body:
            continue
        label = next((key for key in specs if p.text.strip().startswith(key)), None)
        if label is None:
            continue
        p.paragraph_format.keep_with_next = True
        src = doc.add_paragraph()
        src.alignment = WD_ALIGN_PARAGRAPH.CENTER
        src.paragraph_format.space_before = Pt(0)
        src.paragraph_format.space_after = Pt(7)
        run = src.add_run(f"Fuente: {specs[label]}")
        run.font.name = "Arial"
        run.font.size = Pt(9)
        p._p.addnext(src._p)


def add_reproducibility_annex(doc) -> None:
    anchor = find_paragraph(doc, "Anexo PRINCIPAL: CD")
    insert_paragraph_before(anchor, "Anexo 8: Matriz de evidencia técnica y reproducibilidad", "Heading 3")
    insert_paragraph_before(
        anchor,
        "La correspondencia entre objetivos y artefactos verificables es la siguiente: OE1, notebooks/Obj1_Recoleccion_Limpieza.ipynb y src/data_loader.py; OE2, notebooks/Obj2a_Analisis_Patrones.ipynb; OE3, notebooks/Obj3a_Entrenamiento_Arboles_RF.ipynb, notebooks/Obj3b_Entrenamiento_Redes_Neuronales.ipynb y scripts/train_models.py; OE4, notebooks/Obj4_Evaluacion_Segregacion_Riesgo.ipynb, resultados_modelos/metricas_modelos.json y src/predictor.py; OE5, Obj5_Prototipo_Dashboard_Docente/app.py y sus componentes de interfaz.",
    )
    insert_paragraph_before(
        anchor,
        "La verificación mínima consiste en ejecutar el entrenamiento reproducible, revisar el archivo JSON de métricas y correr la suite tests/test_project.py. Los datos identificables no forman parte de la evidencia publicable; el protocolo de privacidad y la ficha del modelo se encuentran en documentacion/08_Protocolo_Privacidad.md y documentacion/09_Ficha_Modelo.md.",
    )


def update_static_lists(doc) -> None:
    toc_anchor = next(p for p in doc.paragraphs if p.text.startswith("4.\tAnálisis de Resultados"))
    insert_list_entries(
        toc_anchor,
        [
            "3.10.\tTrazabilidad metodológica y evidencia por objetivo específico\t—",
            "3.10.1.\tEvidencia del objetivo específico 1: calidad y consolidación\t—",
            "3.10.2.\tEvidencia del objetivo específico 2: análisis de patrones\t—",
            "3.10.3.\tEvidencia del objetivo específico 3: modelado temporal\t—",
            "3.10.4.\tEvidencia del objetivo específico 4: evaluación y segregación\t—",
            "3.10.5.\tEvidencia del objetivo específico 5: prototipo docente\t—",
        ],
    )

    fig_anchor = find_paragraph(doc, "Lista de tablas")
    insert_list_entries(
        fig_anchor,
        [
            "Figura 3.7. Validación, construcción de variables y resultado agregado del objetivo 1\t—",
            "Figura 3.8. Código reproducible y tasas descriptivas de reprobación por asignatura\t—",
            "Figura 3.9. Entrenamiento reproducible y evaluación temporal de los tres clasificadores\t—",
            "Figura 3.10. Trazabilidad de la segregación operativa y distribución de categorías\t—",
            "Figura 3.11. Panel principal del prototipo con filtros e indicadores agregados\t—",
            "Figura 3.12. Simulador libre y separación de salidas\t—",
        ],
    )

    table_anchor = find_paragraph(doc, "Introducción")
    insert_list_entries(
        table_anchor,
        [
            "Tabla 3.1. Matriz de trazabilidad entre objetivos, procesos y evidencias\t—",
            "Tabla 3.2. Reglas de validación y tratamiento de datos académicos\t—",
            "Tabla 3.3. Construcción y partición de la muestra predictiva longitudinal\t—",
        ],
    )


def improve_bibliography_format(doc) -> None:
    paragraphs = doc.paragraphs
    start = next(i for i, p in enumerate(paragraphs) if p.text.strip() == "Bibliografía")
    end = next(i for i, p in enumerate(paragraphs) if p.text.strip() == "Anexos")
    for p in paragraphs[start + 1 : end]:
        if not p.text.strip():
            continue
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        for run in p.runs:
            run.font.name = "Arial"
            run.font.size = Pt(11)


def main() -> None:
    doc = Document(SOURCE)
    normalize_existing_figure_captions(doc)
    add_existing_table_captions(doc)
    add_existing_result_figure_sources(doc)
    add_methodology_section(doc)
    add_reproducibility_annex(doc)
    update_static_lists(doc)
    improve_bibliography_format(doc)
    doc.core_properties.title = "Sistema de alertas tempranas para la prevención del rezago escolar"
    doc.core_properties.subject = "Proyecto final del Diplomado en Ciencia de Datos"
    doc.core_properties.keywords = "rezago escolar, alertas tempranas, ciencia de datos, machine learning"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
